#!/usr/bin/env python3
"""
book_condenser.py

Create a full extractive abridgement of a rights-cleared/public-domain nonfiction book.
The LLM selects paragraph identifiers and writes short analytical metadata only.
All quotation text in the final abridgement is retrieved verbatim from the source.

Supported input formats:
    .epub, .pdf, .docx, .txt, .md

Outputs:
    <output_dir>/
        book_metadata.json
        book_paragraphs.jsonl
        structural_overview.json
        chapter_analysis/*.json
        chapter_analyses.json
        analytical_map.json
        chapter_candidates/*.json
        scored_candidates.json
        global_selection.json
        quality_control.json
        editorial_transitions.json
        editorial_transition_validation.json
        analytical_reading_guide.md
        selection_audit.md
        reading_abridgement.md
        reading_abridgement.pdf
        reading_abridgement.docx       (unless --no-docx)

Example:
    export OPENAI_API_KEY="..."
    book-condenser book.epub \
        --output-dir out/project_maven \
        --target-ratio 0.25 \
        --emphasis "Central thesis, major claims, evidence, concepts, mechanisms, counterarguments, and conclusions."

Design:
    1. Parse the source and assign stable paragraph IDs.
    2. Perform an inspectional structural overview from the book architecture and framing material.
    3. Analyse each chapter for terms, propositions, support, qualifications, and contribution to the whole.
    4. Synthesize an Adlerian analytical map and mandatory reading path.
    5. Nominate contiguous quotation blocks tagged against analytical requirements.
    6. Select blocks under budget while protecting essential propositions and their support.
    7. Check analytical completeness, coherence, and redundancy.
    8. Generate and validate brief, disclosed editorial transitions where omission creates a discontinuity.
    9. Assemble exact source quotations and marked editorial transitions into Markdown, PDF, and DOCX; write a separate analytical guide.
"""

from __future__ import annotations

import argparse
import json
import logging
import math
import os
import posixpath
import re
import shutil
import time
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter
from collections.abc import Sequence
from dataclasses import asdict, dataclass, field
from datetime import datetime
from html import unescape
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Literal, TypeVar

from pydantic import BaseModel, ConfigDict, Field

LOGGER = logging.getLogger("nonfiction_extractive_condenser")
T = TypeVar("T", bound=BaseModel)

WORD_RE = re.compile(r"\b[\w’'-]+\b", re.UNICODE)
HEADING_RE = re.compile(
    r"^\s*(chapter\s+([0-9ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)\b.*|\d+\s+[A-Z][^\n]{2,100}|summary of the framework.*|"
    r"introduction\b.*|prologue\b.*|preface\b.*|foreword\b.*|"
    r"epilogue\b.*|conclusion\b.*|afterword\b.*|"
    r"part\s+([0-9ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)\b.*|"
    r"appendix\b.*|notes\b.*|bibliography\b.*|references\b.*|acknowledg(e)?ments\b.*|index\b.*)$",
    flags=re.IGNORECASE,
)

FRONT_PATTERNS = ("introduction", "prologue", "preface", "foreword")
END_PATTERNS = ("epilogue", "conclusion", "afterword")
EXCLUDE_PATTERNS = (
    "copyright", "contents", "table of contents", "bibliography", "references", "index",
    "notes", "acknowledg", "list of interviews", "written correspondence", "interviews and written",
    "photo insert", "photo-insert", "illustrations", "plates", "image credits", "about the author"
)


# ---------------------------------------------------------------------------
# Local source models
# ---------------------------------------------------------------------------

@dataclass
class Paragraph:
    paragraph_id: str
    chapter_id: str
    chapter_title: str
    index: int
    text: str
    word_count: int
    page: int | None = None


@dataclass
class Chapter:
    chapter_id: str
    title: str
    paragraph_ids: list[str]
    word_count: int
    kind: str = "body"


@dataclass
class Book:
    title: str
    source_path: str
    chapters: list[Chapter]
    paragraphs: dict[str, Paragraph]
    total_words: int
    diagnostics: dict[str, Any] = field(default_factory=dict)


@dataclass
class SelectedBlock:
    block_id: str
    chapter_id: str
    chapter_title: str
    paragraph_ids: list[str]
    word_count: int
    selection_reason: str
    importance: str
    themes: list[str]
    redundancy_risk: str
    block_function: str = "supporting"
    text: str = ""
    score: float = 0.0
    redundant_with: list[str] = field(default_factory=list)
    covers_requirement_ids: list[str] = field(default_factory=list)
    establishes_term_ids: list[str] = field(default_factory=list)
    supports_proposition_ids: list[str] = field(default_factory=list)
    requires_prior_requirement_ids: list[str] = field(default_factory=list)
    protected_requirement_ids: list[str] = field(default_factory=list)
    protected_anchor: bool = False


# ---------------------------------------------------------------------------
# Structured LLM response models
# ---------------------------------------------------------------------------

class StrictModel(BaseModel):
    model_config = ConfigDict(extra="forbid")


class CoreClaim(StrictModel):
    claim: str
    role_in_argument: str
    importance: str


class KeyConcept(StrictModel):
    name: str
    function_in_argument: str
    importance: str


class EvidenceOrCase(StrictModel):
    description: str
    evidentiary_role: str
    importance: str


class ChapterPriority(StrictModel):
    chapter_id: str
    priority: str
    selection_focus: str
    suggested_budget_fraction: float


class StructuralOverview(StrictModel):
    nonfiction_form: Literal[
        "argumentative",
        "historical_investigative_narrative",
        "biography_memoir",
        "technical_explanatory",
        "case_based_policy_business",
        "mixed",
    ]
    central_question: str
    governing_thesis: str
    overview: str
    chronological_or_logical_arc: list[str]
    core_claims: list[CoreClaim]
    key_concepts: list[KeyConcept]
    evidence_and_cases: list[EvidenceOrCase]
    counterarguments_or_limitations: list[str]
    chapter_priorities: list[ChapterPriority]
    selection_rules: list[str]


class ChapterAnalyticalMap(StrictModel):
    analysis_id: str
    chapter_id: str
    chapter_title: str
    chunk_index: int
    chapter_role: Literal[
        "problem_setup", "definition", "claim_development", "mechanism", "evidence",
        "case_study", "objection", "qualification", "synthesis", "conclusion", "mixed"
    ]
    questions_addressed: list[str]
    essential_terms_introduced: list[str]
    propositions_advanced: list[str]
    arguments_or_evidence: list[str]
    limitations_or_qualifications: list[str]
    dependencies_on_previous_chapters: list[str]
    contribution_to_book_unity: str


class AnalyticalRequirement(StrictModel):
    requirement_id: str
    kind: Literal[
        "central_question", "unity_thesis", "essential_term", "major_proposition",
        "supporting_argument", "evidence", "counterargument", "limitation",
        "conclusion", "implication"
    ]
    description: str
    importance: Literal["essential", "important", "supporting"]
    related_claim_ids: list[str]
    preferred_chapter_ids: list[str]
    must_be_preserved: bool


class ArgumentRelation(StrictModel):
    source_requirement_id: str
    target_requirement_id: str
    relation: Literal["defines", "supports", "qualifies", "objects_to", "responds_to", "concludes_from"]


class AnalyticalMap(StrictModel):
    book_classification: str
    central_problem: str
    unity_statement: str
    requirements: list[AnalyticalRequirement]
    argument_relations: list[ArgumentRelation]
    minimum_complete_reading_path: list[str]


class CandidateBlockResponse(StrictModel):
    block_id: str
    paragraph_ids: list[str]
    block_function: Literal[
        "setup", "definition", "claim", "mechanism", "evidence", "representative_episode",
        "turning_point", "consequence", "counterargument", "interpretation", "conclusion"
    ]
    covers_requirement_ids: list[str]
    establishes_term_ids: list[str]
    supports_proposition_ids: list[str]
    requires_prior_requirement_ids: list[str]
    selection_reason: str
    importance: str
    themes: list[str]
    redundancy_risk: str


class ChapterCandidateResponse(StrictModel):
    chapter_id: str
    chapter_contribution: str
    candidate_blocks: list[CandidateBlockResponse]
    omitted_material_description: list[str]


class ScoredBlockResponse(StrictModel):
    block_id: str
    chronological_or_logical_necessity: int = Field(ge=0, le=5)
    institutional_causal_or_argumentative_importance: int = Field(ge=0, le=5)
    turning_point_or_conclusion_value: int = Field(ge=0, le=5)
    explanatory_density: int = Field(ge=0, le=5)
    unity_value: int = Field(ge=0, le=5)
    term_definition_value: int = Field(ge=0, le=5)
    proposition_value: int = Field(ge=0, le=5)
    argument_support_value: int = Field(ge=0, le=5)
    objection_or_limitation_value: int = Field(ge=0, le=5)
    conclusion_value: int = Field(ge=0, le=5)
    analytical_dependency_value: int = Field(ge=0, le=5)
    readability: int = Field(ge=0, le=3)
    redundancy_penalty: int = Field(ge=0, le=5)
    excessive_detail_penalty: int = Field(ge=0, le=5)
    context_dependence_penalty: int = Field(ge=0, le=3)
    redundant_with: list[str]
    rationale: str


class ScoringResponse(StrictModel):
    scores: list[ScoredBlockResponse]


class TransitionNote(StrictModel):
    before_block_id: str
    note: str


class AnalyticalCoverageAssessment(StrictModel):
    unity_preserved: bool
    essential_terms_preserved: bool
    major_propositions_preserved: bool
    supporting_arguments_preserved: bool
    objections_or_limitations_preserved: bool
    conclusion_preserved: bool
    missing_requirement_ids: list[str]
    unsupported_proposition_ids: list[str]
    undefined_term_ids: list[str]


class QualityControlResponse(StrictModel):
    assessment: str
    analytical_coverage: AnalyticalCoverageAssessment
    missing_coverage: list[str]
    remove_block_ids: list[str]
    add_block_ids: list[str]
    transition_notes: list[TransitionNote]


class EditorialTransition(StrictModel):
    before_block_id: str
    after_block_id: str
    transition_type: Literal[
        "omission_bridge", "logical_bridge", "chronological_bridge", "terminological_reminder"
    ]
    text: str
    necessary: bool
    grounding_block_ids: list[str]
    grounding_requirement_ids: list[str]
    rationale: str


class EditorialTransitionResponse(StrictModel):
    transitions: list[EditorialTransition]


class TransitionValidation(StrictModel):
    before_block_id: str
    after_block_id: str
    grounded: bool
    introduces_new_claim: bool
    impersonates_author: bool
    overly_explanatory: bool
    approved_text: str
    reason: str


class TransitionValidationResponse(StrictModel):
    validations: list[TransitionValidation]


# ---------------------------------------------------------------------------
# General utilities
# ---------------------------------------------------------------------------

def word_count(text: str) -> int:
    return len(WORD_RE.findall(text))


def normalize_text(text: str) -> str:
    """Normalize extraction artefacts without rewriting the author's substantive prose."""
    text = text.replace("\u00ad", "").replace("\xa0", " ")
    text = text.replace("\ufb01", "fi").replace("\ufb02", "fl")
    # PDF extraction often represents a broken hyphen with a replacement/noncharacter glyph.
    text = re.sub(r"(?<=\w)[\ufffd\ufffe](?=\w)", "-", text)
    text = text.replace("\ufffd", "").replace("\ufffe", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    # Repair specific drop-cap/OCR splits observed at paragraph openings without
    # collapsing legitimate prose such as "A condensed..." or "I think...".
    text = re.sub(r"^T\s+he\b", "The", text)
    text = re.sub(r"^D\s+uring\b", "During", text)
    text = re.sub(r"\ba(?=[A-Z][a-z]{3,}\b)", "a ", text)
    text = re.sub(r"(?<=[A-Z&])(?=Dr\.)", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text.strip()


def paragraph_split(text: str) -> list[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = re.split(r"\n\s*\n+", text)
    cleaned: list[str] = []
    for block in blocks:
        block = normalize_text(re.sub(r"\n+", " ", block))
        if block and word_count(block) >= 2:
            cleaned.append(block)
    return cleaned


def slug_id(prefix: str, index: int) -> str:
    return f"{prefix}{index:03d}"


def output_dir_slug(source_stem: str) -> str:
    """Build a filesystem-safe folder slug from the source book filename."""
    slug = re.sub(r"[^\w.-]+", "-", source_stem.strip()).strip("-._")
    slug = re.sub(r"-{2,}", "-", slug)
    return slug[:80] or "book"


def allocate_output_dir(
    parent: Path,
    source: Path,
    *,
    reuse: bool = False,
    now: datetime | None = None,
) -> Path:
    """Return a run directory under parent, using a unique name unless reuse is requested."""
    parent = parent.expanduser().resolve()
    if reuse:
        parent.mkdir(parents=True, exist_ok=True)
        return parent

    stamp = (now or datetime.now()).strftime("%Y%m%d-%H%M%S")
    base = f"{output_dir_slug(source.stem)}-{stamp}"
    parent.mkdir(parents=True, exist_ok=True)
    candidate = parent / base
    suffix = 2
    while candidate.exists():
        candidate = parent / f"{base}-{suffix}"
        suffix += 1
    candidate.mkdir(parents=False, exist_ok=False)
    return candidate


def canonical_title(title: str) -> str:
    title = normalize_text(title).lower()
    title = re.sub(r"^[\W_]+|[\W_]+$", "", title)
    title = re.sub(r"\s+", " ", title)
    return title


def classify_chapter(title: str) -> str:
    """Classify reader-visible section titles only; internal filenames are not headings."""
    lower = canonical_title(title)
    excluded_patterns = (
        r"^(copyright|contents|table of contents)\b",
        r"^(bibliography|references|works cited)\b",
        r"^(index)\b",
        r"^(notes|endnotes)\b",
        r"^(acknowledg(e)?ments?)\b",
        r"^(list of interviews|written correspondence|interviews and written)\b",
        r"^(photo[\s-]?insert|illustrations|plates|image credits)\b",
        r"^(about the author)\b",
        r"^(unmapped front matter)\b",
    )
    if any(re.match(pattern, lower) for pattern in excluded_patterns):
        return "exclude"
    if any(re.match(rf"^{re.escape(p)}\b", lower) for p in FRONT_PATTERNS):
        return "front"
    if any(re.match(rf"^{re.escape(p)}\b", lower) for p in END_PATTERNS):
        return "ending"
    return "body"


def safe_json_dump(data: Any, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = model_dump_jsonable(data)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def model_dump_jsonable(obj: Any) -> Any:
    if isinstance(obj, BaseModel):
        return obj.model_dump()
    if hasattr(obj, "__dataclass_fields__"):
        return asdict(obj)
    if isinstance(obj, list):
        return [model_dump_jsonable(x) for x in obj]
    if isinstance(obj, dict):
        return {k: model_dump_jsonable(v) for k, v in obj.items()}
    return obj


def truncate_words(text: str, max_words: int) -> str:
    words = text.split()
    if len(words) <= max_words:
        return text
    front = max_words * 2 // 3
    back = max_words - front
    return " ".join(words[:front]) + "\n\n[... omitted for input-length control ...]\n\n" + " ".join(words[-back:])


# ---------------------------------------------------------------------------
# Source loading
# ---------------------------------------------------------------------------

def build_book_from_sections(
    source_path: Path,
    title: str,
    sections: list[tuple[str, list[tuple[str, int | None]]]],
    diagnostics: dict[str, Any] | None = None,
) -> Book:
    paragraphs: dict[str, Paragraph] = {}
    chapters: list[Chapter] = []

    usable_sections = [(t.strip() or f"Section {i + 1}", ps) for i, (t, ps) in enumerate(sections) if ps]
    if not usable_sections:
        raise ValueError("No readable paragraph text could be extracted from the input file.")

    for chapter_no, (chapter_title, blocks) in enumerate(usable_sections, start=1):
        chapter_id = slug_id("CH", chapter_no)
        para_ids: list[str] = []
        chapter_words = 0
        for para_no, (text, page) in enumerate(blocks, start=1):
            text = normalize_text(text)
            if not text or word_count(text) < 2:
                continue
            pid = f"{chapter_id}-P{para_no:04d}"
            wc = word_count(text)
            paragraphs[pid] = Paragraph(
                paragraph_id=pid,
                chapter_id=chapter_id,
                chapter_title=chapter_title,
                index=para_no,
                text=text,
                word_count=wc,
                page=page,
            )
            para_ids.append(pid)
            chapter_words += wc
        if para_ids:
            chapter_kind = classify_chapter(chapter_title)
            # EPUBs often contain a separate PART title page with no substantive prose.
            # Preserve it in the parse audit but never allocate quotation budget to it.
            if re.match(r"^part\b", canonical_title(chapter_title)) and chapter_words < 50:
                chapter_kind = "exclude"
            chapters.append(
                Chapter(
                    chapter_id=chapter_id,
                    title=chapter_title,
                    paragraph_ids=para_ids,
                    word_count=chapter_words,
                    kind=chapter_kind,
                )
            )

    return Book(
        title=title,
        source_path=str(source_path),
        chapters=chapters,
        paragraphs=paragraphs,
        total_words=sum(c.word_count for c in chapters if c.kind != "exclude"),
        diagnostics=diagnostics or {},
    )


def load_txt_or_md(path: Path) -> Book:
    raw = path.read_text(encoding="utf-8", errors="replace")
    lines = raw.splitlines()
    sections: list[tuple[str, list[tuple[str, int | None]]]] = []
    current_title = "Opening"
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        text = "\n".join(current_lines).strip()
        if text:
            sections.append((current_title, [(p, None) for p in paragraph_split(text)]))
        current_lines = []

    for line in lines:
        stripped = line.strip().lstrip("#").strip()
        is_markdown_heading = bool(re.match(r"^\s{0,3}#{1,6}\s+\S", line))
        is_heading = is_markdown_heading or (len(stripped) < 120 and HEADING_RE.match(stripped))
        if is_heading:
            flush()
            current_title = stripped
        else:
            current_lines.append(line)
    flush()

    if not sections:
        sections = [("Book text", [(p, None) for p in paragraph_split(raw)])]
    return build_book_from_sections(path, path.stem, sections)


def load_docx(path: Path) -> Book:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("Install python-docx to read DOCX files: pip install python-docx") from exc

    doc = Document(str(path))
    sections: list[tuple[str, list[tuple[str, int | None]]]] = []
    current_title = "Opening"
    current: list[tuple[str, int | None]] = []

    def flush() -> None:
        nonlocal current
        if current:
            sections.append((current_title, current))
            current = []

    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue
        style_name = getattr(para.style, "name", "") or ""
        is_heading = style_name.lower().startswith("heading") or (
            len(text) < 120 and HEADING_RE.match(text)
        )
        if is_heading:
            flush()
            current_title = text
        else:
            current.append((text, None))
    flush()
    return build_book_from_sections(path, path.stem, sections)


class _EPUBBlockParser(HTMLParser):
    """Forgiving EPUB content reader with wrapper-anchor propagation."""
    TEXT_TAGS = {"p", "blockquote", "li", "td"}
    HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
    WRAPPER_TAGS = {"section", "article", "div", "aside", "main", "span", "a"}
    SKIP_TAGS = {"script", "style", "nav", "svg"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.skip_depth = 0
        self.active_tag: str | None = None
        self.active_ids: set[str] = set()
        self.pending_ids: set[str] = set()
        self.buffer: list[str] = []
        self.blocks: list[tuple[str, str, set[str]]] = []
        self.semantic_tokens: set[str] = set()
        self.image_count = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        attr = {k.lower(): (v or "") for k, v in attrs}
        identifiers = {attr[key] for key in ("id", "name") if attr.get(key)}
        for key, value in attr.items():
            if key == "role" or key.endswith("}type") or key in {"epub:type", "type"}:
                self.semantic_tokens.update(value.lower().replace(":", "-").split())
        if tag == "img":
            self.image_count += 1
        if tag in self.SKIP_TAGS:
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        if tag in self.WRAPPER_TAGS and identifiers:
            self.pending_ids.update(identifiers)
        if tag in self.TEXT_TAGS | self.HEADING_TAGS:
            self._flush()
            self.active_tag = tag
            self.active_ids = identifiers | self.pending_ids
            self.pending_ids.clear()
            self.buffer = []
        elif tag == "br" and self.active_tag:
            self.buffer.append(" ")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS:
            self.skip_depth = max(0, self.skip_depth - 1)
            return
        if not self.skip_depth and self.active_tag == tag:
            self._flush()

    def handle_data(self, data: str) -> None:
        if not self.skip_depth and self.active_tag:
            self.buffer.append(data)

    def close(self) -> None:
        super().close()
        self._flush()

    def _flush(self) -> None:
        if self.active_tag:
            value = normalize_text(unescape(" ".join(self.buffer)))
            if value:
                self.blocks.append((self.active_tag, value, set(self.active_ids)))
        self.active_tag = None
        self.active_ids = set()
        self.buffer = []


class _LooseEPUBTextParser(HTMLParser):
    """Fallback preserving order when prose is carried directly by layout containers."""
    BREAK_TAGS = {"p", "blockquote", "li", "td", "div", "section", "article"}
    HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
    WRAPPER_TAGS = {"div", "section", "article", "span", "a"}
    SKIP_TAGS = {"script", "style", "nav", "svg"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.skip_depth = 0
        self.heading_tag: str | None = None
        self.heading_buffer: list[str] = []
        self.prose_buffer: list[str] = []
        self.pending_ids: set[str] = set()
        self.events: list[tuple[str, str, set[str]]] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        attr = {k.lower(): (v or "") for k, v in attrs}
        ids = {attr[key] for key in ("id", "name") if attr.get(key)}
        if tag in self.SKIP_TAGS:
            self._flush_prose()
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        if tag in self.WRAPPER_TAGS and ids:
            self.pending_ids.update(ids)
        if tag in self.HEADING_TAGS:
            self._flush_prose()
            self.heading_tag = tag
            self.heading_buffer = []
            self.pending_ids.update(ids)
        elif tag in self.BREAK_TAGS and not self.heading_tag:
            self._flush_prose()
            self.pending_ids.update(ids)
        elif tag == "br" and not self.heading_tag:
            self.prose_buffer.append(" ")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS:
            self.skip_depth = max(0, self.skip_depth - 1)
            return
        if self.skip_depth:
            return
        if self.heading_tag == tag:
            value = normalize_text(unescape(" ".join(self.heading_buffer)))
            if value:
                self.events.append((tag, value, set(self.pending_ids)))
            self.pending_ids.clear()
            self.heading_tag = None
            self.heading_buffer = []
        elif tag in self.BREAK_TAGS:
            self._flush_prose()

    def handle_data(self, data: str) -> None:
        if self.skip_depth:
            return
        if self.heading_tag:
            self.heading_buffer.append(data)
        else:
            self.prose_buffer.append(data)

    def close(self) -> None:
        super().close()
        self._flush_prose()

    def _flush_prose(self) -> None:
        value = normalize_text(unescape(" ".join(self.prose_buffer)))
        if value and word_count(value) >= 2:
            self.events.append(("p", value, set(self.pending_ids)))
            self.pending_ids.clear()
        self.prose_buffer = []


def _zip_read_text(archive: zipfile.ZipFile, member: str) -> str:
    return archive.read(member).decode("utf-8", errors="replace")


def _epub_opf_path(archive: zipfile.ZipFile) -> str:
    try:
        root = ET.fromstring(_zip_read_text(archive, "META-INF/container.xml"))
        for element in root.iter():
            if element.tag.rsplit("}", 1)[-1] == "rootfile":
                candidate = element.attrib.get("full-path")
                if candidate:
                    return candidate
    except (KeyError, ET.ParseError):
        pass
    opfs = [name for name in archive.namelist() if name.lower().endswith(".opf")]
    if not opfs:
        raise ValueError("EPUB does not contain an OPF package document.")
    return opfs[0]


def _epub_package_data(archive: zipfile.ZipFile) -> dict[str, Any]:
    opf_path = _epub_opf_path(archive)
    opf_dir = posixpath.dirname(opf_path)
    root = ET.fromstring(_zip_read_text(archive, opf_path))
    title = ""
    fixed_layout_indicators: list[str] = []
    manifest: dict[str, dict[str, str]] = {}

    for element in root.iter():
        local = element.tag.rsplit("}", 1)[-1]
        if local == "title" and element.text and not title:
            title = normalize_text(element.text)
        elif local == "meta":
            prop = element.attrib.get("property", "").lower()
            name = element.attrib.get("name", "").lower()
            value = normalize_text(element.text or element.attrib.get("content", "")).lower()
            if ("rendition:layout" in prop or "fixed-layout" in name) and "pre-paginated" in value:
                fixed_layout_indicators.append("OPF metadata declares pre-paginated/fixed layout")
        elif local == "item":
            item_id = element.attrib.get("id")
            href = element.attrib.get("href")
            if item_id and href:
                manifest[item_id] = {
                    "href": posixpath.normpath(posixpath.join(opf_dir, href)),
                    "media_type": element.attrib.get("media-type", ""),
                    "properties": element.attrib.get("properties", ""),
                }

    spine: list[str] = []
    toc_id: str | None = None
    for element in root.iter():
        local = element.tag.rsplit("}", 1)[-1]
        if local == "spine":
            toc_id = element.attrib.get("toc")
            if "pre-paginated" in element.attrib.get("properties", "").lower():
                fixed_layout_indicators.append("Spine declares pre-paginated layout")
        elif local == "itemref":
            idref = element.attrib.get("idref")
            if idref in manifest:
                spine.append(manifest[idref]["href"])

    ncx_path: str | None = manifest.get(toc_id or "", {}).get("href")
    nav_path: str | None = None
    content_docs: list[str] = []
    image_count = 0
    for item in manifest.values():
        media_type = item["media_type"].lower()
        props = set(item["properties"].lower().split())
        if "nav" in props:
            nav_path = item["href"]
        if "html" in media_type or "xhtml" in media_type:
            content_docs.append(item["href"])
        if media_type.startswith("image/"):
            image_count += 1
        if "svg" in media_type:
            fixed_layout_indicators.append("Manifest contains SVG resources")

    return {
        "title": title or "Untitled book",
        "spine": spine,
        "ncx_path": ncx_path,
        "nav_path": nav_path,
        "content_docs": content_docs,
        "image_count": image_count,
        "fixed_layout_indicators": sorted(set(fixed_layout_indicators)),
    }


def _join_href(base: str, src: str) -> tuple[str, str | None]:
    path_part, _, anchor = src.partition("#")
    return posixpath.normpath(posixpath.join(base, path_part)), (anchor or None)


def _append_nav_entry(
    mapping: dict[str, list[tuple[str | None, str]]],
    base: str,
    src: str,
    label: str,
    excluded_scope: str | None = None,
) -> None:
    if not src or not label:
        return
    href, anchor = _join_href(base, src)
    emitted = f"{excluded_scope}: {label}" if excluded_scope and classify_chapter(label) != "exclude" else label
    mapping.setdefault(href, []).append((anchor, emitted))


def _epub_ncx_map(archive: zipfile.ZipFile, toc_path: str | None) -> dict[str, list[tuple[str | None, str]]]:
    mapping: dict[str, list[tuple[str | None, str]]] = {}
    if not toc_path:
        return mapping
    try:
        root = ET.fromstring(_zip_read_text(archive, toc_path))
    except (KeyError, ET.ParseError):
        return mapping
    base = posixpath.dirname(toc_path)

    def direct_child(element: ET.Element, name: str) -> ET.Element | None:
        return next((x for x in list(element) if x.tag.rsplit("}", 1)[-1] == name), None)

    def walk(navpoint: ET.Element, excluded_scope: str | None = None) -> None:
        nav_label = direct_child(navpoint, "navLabel")
        label = ""
        if nav_label is not None:
            text_node = next((x for x in nav_label.iter() if x.tag.rsplit("}", 1)[-1] == "text"), None)
            label = normalize_text(text_node.text or "") if text_node is not None else ""
        content = direct_child(navpoint, "content")
        src = content.attrib.get("src", "") if content is not None else ""
        _append_nav_entry(mapping, base, src, label, excluded_scope)
        next_scope = label if label and classify_chapter(label) == "exclude" else excluded_scope
        for child in list(navpoint):
            if child.tag.rsplit("}", 1)[-1] == "navPoint":
                walk(child, next_scope)

    nav_map = next((x for x in root.iter() if x.tag.rsplit("}", 1)[-1] == "navMap"), None)
    if nav_map is not None:
        for child in list(nav_map):
            if child.tag.rsplit("}", 1)[-1] == "navPoint":
                walk(child)
    return mapping


def _epub_type(element: ET.Element) -> set[str]:
    values: list[str] = []
    for key, value in element.attrib.items():
        local = key.rsplit("}", 1)[-1].lower()
        if local in {"type", "role"}:
            values.extend((value or "").lower().replace(":", "-").split())
    return set(values)


def _epub3_nav_data(
    archive: zipfile.ZipFile, nav_path: str | None
) -> tuple[dict[str, list[tuple[str | None, str]]], dict[str, str]]:
    toc_map: dict[str, list[tuple[str | None, str]]] = {}
    landmarks: dict[str, str] = {}
    if not nav_path:
        return toc_map, landmarks
    try:
        root = ET.fromstring(_zip_read_text(archive, nav_path))
    except (KeyError, ET.ParseError):
        return toc_map, landmarks
    base = posixpath.dirname(nav_path)
    navs = [x for x in root.iter() if x.tag.rsplit("}", 1)[-1] == "nav"]
    toc_nav = next((n for n in navs if "toc" in _epub_type(n)), None)
    landmarks_nav = next((n for n in navs if "landmarks" in _epub_type(n)), None)

    def walk_li(element: ET.Element, mapping: dict[str, list[tuple[str | None, str]]], excluded_scope: str | None = None) -> None:
        for li in [x for x in list(element) if x.tag.rsplit("}", 1)[-1] == "li"]:
            anchor = next((x for x in list(li) if x.tag.rsplit("}", 1)[-1] == "a"), None)
            label = normalize_text(" ".join(anchor.itertext())) if anchor is not None else ""
            src = anchor.attrib.get("href", "") if anchor is not None else ""
            _append_nav_entry(mapping, base, src, label, excluded_scope)
            next_scope = label if label and classify_chapter(label) == "exclude" else excluded_scope
            for child in list(li):
                if child.tag.rsplit("}", 1)[-1] == "ol":
                    walk_li(child, mapping, next_scope)

    if toc_nav is not None:
        for child in list(toc_nav):
            if child.tag.rsplit("}", 1)[-1] == "ol":
                walk_li(child, toc_map)

    if landmarks_nav is not None:
        for anchor in [x for x in landmarks_nav.iter() if x.tag.rsplit("}", 1)[-1] == "a"]:
            src = anchor.attrib.get("href", "")
            label = normalize_text(" ".join(anchor.itertext()))
            role_tokens = _epub_type(anchor)
            href, _ = _join_href(base, src)
            semantic_label = label
            if role_tokens & {"backmatter", "bibliography", "endnotes", "rearnotes", "index", "acknowledgments", "copyright-page"}:
                semantic_label = label if classify_chapter(label) == "exclude" else f"Notes: {label}"
            if href and semantic_label:
                landmarks[href] = semantic_label
    return toc_map, landmarks


def _epub_semantic_exclusion(tokens: set[str]) -> str | None:
    mapping = {
        "backmatter": "Notes",
        "endnotes": "Notes",
        "rearnotes": "Notes",
        "footnotes": "Notes",
        "bibliography": "Bibliography",
        "index": "Index",
        "acknowledgments": "Acknowledgments",
        "copyright-page": "Copyright",
        "colophon": "Copyright",
    }
    for token, label in mapping.items():
        if token in tokens:
            return label
    return None


def _epub_html_blocks(raw: str) -> tuple[list[tuple[str, str, set[str]]], set[str], int, bool]:
    parser = _EPUBBlockParser()
    parser.feed(raw)
    parser.close()
    primary_words = sum(word_count(value) for tag, value, _ in parser.blocks if tag in parser.TEXT_TAGS)
    used_loose_fallback = False
    blocks = parser.blocks
    if primary_words < 20:
        fallback = _LooseEPUBTextParser()
        fallback.feed(raw)
        fallback.close()
        recovered_words = sum(
            word_count(value) for tag, value, _ in fallback.events if tag == "p"
        )
        if recovered_words > primary_words:
            blocks = fallback.events
            used_loose_fallback = True
    return blocks, parser.semantic_tokens, parser.image_count, used_loose_fallback


def _visible_section_title(blocks: list[tuple[str, str, set[str]]]) -> str | None:
    for tag, value, _ in blocks[:10]:
        if tag in {"h1", "h2", "h3"} and word_count(value) <= 24:
            return value
        if tag in {"p", "li"} and word_count(value) <= 16 and HEADING_RE.match(value):
            return value
    return None


def _split_blocks_by_navigation(
    blocks: list[tuple[str, str, set[str]]],
    toc_entries: list[tuple[str | None, str]],
) -> list[tuple[str, list[tuple[str, int | None]]]]:
    readable_tags = _EPUBBlockParser.TEXT_TAGS
    anchor_to_title = {anchor: label for anchor, label in toc_entries if anchor}
    if len(toc_entries) <= 1 or not anchor_to_title:
        readable = [(value, None) for tag, value, _ in blocks if tag in readable_tags and word_count(value) >= 2]
        return [(toc_entries[0][1], readable)] if readable else []
    sections: list[tuple[str, list[tuple[str, int | None]]]] = []
    current_title = toc_entries[0][1]
    current_blocks: list[tuple[str, int | None]] = []
    for tag, value, ids in blocks:
        matching = next((anchor_to_title[a] for a in ids if a in anchor_to_title), None)
        if matching and current_blocks:
            sections.append((current_title, current_blocks))
            current_blocks = []
            current_title = matching
        elif matching:
            current_title = matching
        if tag in readable_tags and word_count(value) >= 2:
            current_blocks.append((value, None))
    if current_blocks:
        sections.append((current_title, current_blocks))
    return sections


def load_epub(path: Path) -> Book:
    sections: list[tuple[str, list[tuple[str, int | None]]]] = []
    pending_front: list[tuple[str, int | None]] = []
    active_excluded_scope: str | None = None
    fallback_items = 0
    loose_prose_items = 0
    unassigned_words = 0
    total_images_in_content = 0
    warnings: list[str] = []

    with zipfile.ZipFile(path) as archive:
        package = _epub_package_data(archive)
        title = package["title"]
        spine = package["spine"] or package["content_docs"]
        epub3_toc, landmarks = _epub3_nav_data(archive, package["nav_path"])
        ncx_toc = _epub_ncx_map(archive, package["ncx_path"])
        toc_map = epub3_toc or ncx_toc
        navigation_source = "EPUB3 nav.xhtml" if epub3_toc else ("NCX" if ncx_toc else "visible-heading fallback")

        for member in spine:
            if member not in archive.namelist():
                warnings.append(f"Spine resource missing from archive: {member}")
                continue
            blocks, semantic_tokens, image_count, used_loose = _epub_html_blocks(_zip_read_text(archive, member))
            total_images_in_content += image_count
            if used_loose:
                loose_prose_items += 1
            readable = [(value, None) for tag, value, _ in blocks if tag in _EPUBBlockParser.TEXT_TAGS and word_count(value) >= 2]
            if not readable:
                continue

            semantic_label = _epub_semantic_exclusion(semantic_tokens) or landmarks.get(member)
            toc_entries = toc_map.get(member, [])
            if semantic_label and classify_chapter(semantic_label) == "exclude":
                sections.append((semantic_label, readable))
                active_excluded_scope = semantic_label
                continue

            if toc_entries:
                recovered = _split_blocks_by_navigation(blocks, toc_entries)
                sections.extend(recovered)
                active_excluded_scope = (
                    toc_entries[-1][1] if recovered and classify_chapter(toc_entries[-1][1]) == "exclude" else None
                )
                continue

            if active_excluded_scope and sections:
                sections[-1][1].extend(readable)
                continue

            fallback_items += 1
            recovered_sections: list[tuple[str, list[tuple[str, int | None]]]] = []
            current_title: str | None = None
            current_blocks: list[tuple[str, int | None]] = []
            for tag, value, _ in blocks:
                heading_like = (
                    (tag in {"h1", "h2", "h3"} and word_count(value) <= 24)
                    or (tag in {"p", "li"} and word_count(value) <= 16 and HEADING_RE.match(value))
                )
                if heading_like:
                    if current_title and current_blocks:
                        recovered_sections.append((current_title, current_blocks))
                    current_title = value
                    current_blocks = []
                elif tag in _EPUBBlockParser.TEXT_TAGS and word_count(value) >= 2:
                    current_blocks.append((value, None))
            if current_title and current_blocks:
                recovered_sections.append((current_title, current_blocks))

            if recovered_sections:
                sections.extend(recovered_sections)
            elif sections:
                sections[-1][1].extend(readable)
            else:
                pending_front.extend(readable)
                unassigned_words += sum(word_count(value) for value, _ in readable)

        if pending_front:
            sections.insert(0, ("Unmapped front matter", pending_front))

    if not sections:
        raise ValueError("EPUB parsing returned no readable content.")

    extracted_words = sum(word_count(value) for _, blocks in sections for value, _ in blocks)
    fixed_indicators = package["fixed_layout_indicators"]
    image_only_risk = "low"
    if extracted_words < 500 and (package["image_count"] > 5 or total_images_in_content > 5):
        image_only_risk = "high"
    elif fixed_indicators or (total_images_in_content > len(spine) and extracted_words < 5000):
        image_only_risk = "medium"
    if loose_prose_items:
        warnings.append(f"Recovered prose from nonstandard container markup in {loose_prose_items} spine item(s).")
    if fallback_items:
        warnings.append(f"Used visible-heading/continuation recovery in {fallback_items} spine item(s).")
    if fixed_indicators:
        warnings.append("Fixed-layout/SVG indicators detected; verify the structure report and sample text.")
    proceed = image_only_risk != "high" and extracted_words >= 500
    confidence = "high"
    if navigation_source == "visible-heading fallback" or loose_prose_items or image_only_risk == "medium":
        confidence = "medium"
    if not proceed:
        confidence = "low"

    diagnostics = {
        "format": "EPUB",
        "navigation_source": navigation_source,
        "spine_items_found": len(spine),
        "sections_recovered_before_classification": len(sections),
        "unassigned_text_words": unassigned_words,
        "prose_fallback_items": loose_prose_items,
        "fixed_layout_indicators": fixed_indicators,
        "manifest_image_resources": package["image_count"],
        "content_image_elements": total_images_in_content,
        "image_only_content_risk": image_only_risk,
        "parser_confidence": confidence,
        "proceed_with_llm_extraction": proceed,
        "warnings": warnings,
    }
    return build_book_from_sections(path, title, sections, diagnostics=diagnostics)


PDF_BOILERPLATE_RE = re.compile(
    r"(Downloaded from\s+https?://academic\.oup\.com/\S+.*?$|"
    r"The Edge of Sentience:.*?Oxford University Press\..*?$|"
    r"©\s*Jonathan Birch.*?$|"
    r"DOI:\s*\S+.*?$)",
    flags=re.IGNORECASE | re.MULTILINE,
)
PDF_MAIN_TITLE_RE = re.compile(
    r"^(summary of the framework.*|introduction|preface|prologue|foreword|"
    r"\d+\s+\S.*|conclusion|epilogue|afterword|bibliography|references|index|"\
    r"list of interviews.*|written correspondence.*|photo[- ]insert.*|illustrations|plates)$",
    flags=re.IGNORECASE,
)


def normalize_pdf_line(line: str) -> str:
    line = line.replace("\u00ad", "").replace("\ufb01", "fi").replace("\ufb02", "fl")
    line = line.replace("\ufffd", "")
    line = re.sub(r"\s+", " ", line).strip()
    return line


def normalize_repeated_line(line: str) -> str:
    line = normalize_pdf_line(line).lower()
    line = re.sub(r"\d+", "#", line)
    return re.sub(r"[^a-z# ]+", "", line).strip()


def pdf_margin_noise_lines(doc: Any) -> set[str]:
    """Identify short running headers/footers recurring across PDF pages."""
    counts: Counter[str] = Counter()
    for page in doc:
        lines = [normalize_pdf_line(x) for x in page.get_text("text", sort=True).splitlines()]
        marginal = lines[:3] + lines[-4:]
        for line in marginal:
            key = normalize_repeated_line(line)
            if key and len(key) <= 100:
                counts[key] += 1
    threshold = max(3, int(len(doc) * 0.06))
    return {line for line, count in counts.items() if count >= threshold}


def clean_pdf_block(raw: str, recurring_noise: set[str]) -> str:
    raw = raw.replace("\u00ad", "").replace("\ufb01", "fi").replace("\ufb02", "fl")
    raw = re.sub(r"(?<=\w)[\ufffd\ufffe](?=\w)", "-", raw)
    raw = raw.replace("\ufffd", "").replace("\ufffe", "")
    raw = PDF_BOILERPLATE_RE.sub("", raw)
    lines: list[str] = []
    for candidate in raw.splitlines():
        line = normalize_pdf_line(candidate)
        if not line:
            continue
        if re.match(r"^\d+$", line):
            continue
        if normalize_repeated_line(line) in recurring_noise:
            continue
        if line.lower().startswith("downloaded from https://academic.oup.com"):
            continue
        lines.append(line)
    if not lines:
        return ""

    # Restore prose split by PDF line wraps while repairing line-break hyphenation.
    text = lines[0]
    for line in lines[1:]:
        if text.endswith("-") and line and line[0].islower():
            text = text[:-1] + line
        else:
            text += " " + line
    text = re.sub(r"(?<=\w)\s*[-‐‑]\s+(?=\w)", "-", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return normalize_text(text)


def load_manual_chapter_map(path: Path) -> list[dict[str, Any]]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, list):
        raise ValueError("--chapter-map must be a JSON list of objects.")
    for entry in data:
        if not all(k in entry for k in ("title", "start_page")):
            raise ValueError("Every chapter-map entry must contain title and start_page.")
    return data


def pdf_outline_sections(doc: Any) -> list[dict[str, Any]]:
    toc = doc.get_toc(simple=True)
    if not toc:
        return []
    candidates: list[dict[str, Any]] = []
    for level, title, page in toc:
        title = normalize_text(title)
        if PDF_MAIN_TITLE_RE.match(title):
            candidates.append({"level": level, "title": title, "start_page": int(page)})
    # Prefer the outline level with a plausible sequence of book chapters.
    by_level: dict[int, list[dict[str, Any]]] = {}
    for item in candidates:
        by_level.setdefault(item["level"], []).append(item)
    plausible = [items for items in by_level.values() if len(items) >= 3]
    chosen = max(plausible, key=len) if plausible else candidates
    chosen = sorted({(x["title"], x["start_page"]): x for x in chosen}.values(), key=lambda x: x["start_page"])
    return chosen



def pdf_detect_sections_from_text(doc: Any, recurring_noise: set[str]) -> list[dict[str, Any]]:
    """Recover chapter starts from visible page text when a PDF outline is absent."""
    found: list[dict[str, Any]] = []
    for page_no, page in enumerate(doc, start=1):
        lines = [normalize_pdf_line(line) for line in page.get_text("text", sort=True).splitlines()]
        candidates = [
            line for line in lines[:12]
            if line and normalize_repeated_line(line) not in recurring_noise and len(line) <= 140
        ]
        for candidate in candidates:
            if PDF_MAIN_TITLE_RE.match(candidate) or HEADING_RE.match(candidate):
                found.append({"title": candidate, "start_page": page_no})
                break
    deduped: list[dict[str, Any]] = []
    seen: set[tuple[str, int]] = set()
    for item in found:
        key = (canonical_title(str(item["title"])), int(item["start_page"]))
        if key not in seen:
            deduped.append(item)
            seen.add(key)
    return deduped if len(deduped) >= 3 else []


def load_pdf(path: Path, chapter_map_path: Path | None = None) -> Book:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError("Install PyMuPDF to read PDFs robustly: pip install pymupdf") from exc

    doc = fitz.open(str(path))
    metadata = doc.metadata or {}
    title = str(metadata.get("title") or path.stem)
    recurring_noise = pdf_margin_noise_lines(doc)
    outline = load_manual_chapter_map(chapter_map_path) if chapter_map_path else pdf_outline_sections(doc)
    if not outline:
        outline = pdf_detect_sections_from_text(doc, recurring_noise)

    if not outline:
        visible_words = sum(word_count(page.get_text("text", sort=True) or "") for page in doc)
        if visible_words < 500:
            raise ValueError(
                "This PDF contains little or no extractable text. Run OCR first, then re-run the condenser."
            )
        raise ValueError(
            "No reliable chapter outline could be recovered from PDF bookmarks or visible headings. "
            "Supply --chapter-map chapters.json. The condenser will not spend API calls on an "
            "undifferentiated full-book extraction."
        )

    starts = [int(x["start_page"]) for x in outline]
    sections: list[tuple[str, list[tuple[str, int | None]]]] = []
    for idx, entry in enumerate(outline):
        start_page = int(entry["start_page"])
        end_page = int(entry.get("end_page") or ((starts[idx + 1] - 1) if idx + 1 < len(starts) else len(doc)))
        title_text = str(entry["title"]).strip()
        blocks: list[tuple[str, int | None]] = []
        for page_no in range(start_page, min(end_page, len(doc)) + 1):
            page = doc[page_no - 1]
            for block in page.get_text("blocks", sort=True):
                cleaned = clean_pdf_block(str(block[4]), recurring_noise)
                if cleaned and word_count(cleaned) >= 3:
                    blocks.append((cleaned, page_no))
        if blocks:
            sections.append((title_text, blocks))

    if not sections:
        raise ValueError("PDF parsing returned no readable chapter content.")
    extracted_words = sum(word_count(value) for _, blocks in sections for value, _ in blocks)
    page_image_count = sum(len(page.get_images(full=True)) for page in doc)
    source = "manual chapter map" if chapter_map_path else ("PDF bookmarks" if pdf_outline_sections(doc) else "visible-heading fallback")
    image_risk = "high" if extracted_words < 500 and page_image_count > len(doc) * 0.5 else ("medium" if page_image_count > len(doc) and extracted_words < 5000 else "low")
    diagnostics = {
        "format": "PDF",
        "navigation_source": source,
        "pages": len(doc),
        "sections_recovered_before_classification": len(sections),
        "unassigned_text_words": 0,
        "fixed_layout_indicators": ["PDF pages are fixed-layout by definition"],
        "content_image_elements": page_image_count,
        "image_only_content_risk": image_risk,
        "parser_confidence": "low" if image_risk == "high" else ("medium" if source == "visible-heading fallback" or image_risk == "medium" else "high"),
        "proceed_with_llm_extraction": image_risk != "high" and extracted_words >= 500,
        "warnings": ["Verify visible-heading recovery before full extraction."] if source == "visible-heading fallback" else [],
    }
    return build_book_from_sections(path, title, sections, diagnostics=diagnostics)


def load_book(path: Path, chapter_map_path: Path | None = None) -> Book:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return load_txt_or_md(path)
    if suffix == ".docx":
        return load_docx(path)
    if suffix == ".epub":
        return load_epub(path)
    if suffix == ".pdf":
        return load_pdf(path, chapter_map_path=chapter_map_path)
    raise ValueError(f"Unsupported input type: {suffix}. Use EPUB, PDF, DOCX, TXT, or MD.")


def validate_book_structure(book: Book) -> None:
    included = [c for c in book.chapters if c.kind != "exclude" and c.word_count > 0]
    diagnostics = book.diagnostics or {}
    if diagnostics.get("proceed_with_llm_extraction") is False:
        raise ValueError(
            "Parser confidence is insufficient for extractive condensation: the source may be fixed-layout, "
            "image-based, or insufficiently recoverable as prose. Review parsed_structure_report.md; "
            "use OCR, a better EPUB export, or a manual PDF chapter map before proceeding."
        )
    excluded_words = sum(c.word_count for c in book.chapters if c.kind == "exclude")
    all_words = sum(c.word_count for c in book.chapters)
    if book.total_words <= 0:
        raise ValueError(
            "No included book text remains after section classification. The source structure was likely "
            "misread or every section was mistaken for back matter. Review parsed_structure_report.md."
        )
    if book.total_words > 30000 and len(included) < 3:
        raise ValueError(
            f"Structural parsing is implausible: {book.total_words:,} words were assigned to only "
            f"{len(included)} included section(s). Provide a chapter map for PDF input or inspect the EPUB navigation."
        )
    if all_words > 0 and excluded_words / all_words > 0.50:
        raise ValueError(
            f"Structural parsing is implausible: {excluded_words / all_words:.1%} of extracted words were "
            "classified as back matter. Review parsed_structure_report.md before proceeding."
        )
    if excluded_words > 0:
        LOGGER.info("Excluded non-reading matter was detected and will not be used for passage selection.")


# ---------------------------------------------------------------------------
# LLM interface
# ---------------------------------------------------------------------------

class LLM:
    def __init__(self, model: str, retries: int = 3) -> None:
        if not os.environ.get("OPENAI_API_KEY"):
            raise RuntimeError(
                "OPENAI_API_KEY is not set. Set it in your environment before running the full pipeline, "
                "or use --parse-only to validate a source without API calls."
            )
        try:
            from openai import OpenAI
        except ImportError as exc:
            raise ImportError(
                "The full pipeline requires the current OpenAI Python SDK. "
                "Install or upgrade it with: pip install --upgrade openai"
            ) from exc
        self.client = OpenAI()
        self.model = model
        self.retries = retries

    def structured(self, instructions: str, user_input: str, response_model: type[T]) -> T:
        last_error: Exception | None = None
        for attempt in range(1, self.retries + 1):
            try:
                response = self.client.responses.parse(
                    model=self.model,
                    input=[
                        {"role": "system", "content": instructions},
                        {"role": "user", "content": user_input},
                    ],
                    text_format=response_model,
                )
                parsed = response.output_parsed
                if parsed is None:
                    raise RuntimeError("The API returned no parsed structured output.")
                return parsed
            except Exception as exc:  # API/network/schema error handling with limited retries
                last_error = exc
                if attempt == self.retries:
                    break
                delay = min(2 ** attempt, 10)
                LOGGER.warning("LLM call failed on attempt %d/%d: %s. Retrying in %ds.", attempt, self.retries, exc, delay)
                time.sleep(delay)
        raise RuntimeError(f"LLM structured-output call failed after {self.retries} attempts: {last_error}") from last_error


# ---------------------------------------------------------------------------
# Prompt formatting
# ---------------------------------------------------------------------------

def render_paragraphs(book: Book, paragraph_ids: Sequence[str], max_words: int | None = None) -> str:
    rendered = []
    used = 0
    for pid in paragraph_ids:
        para = book.paragraphs[pid]
        if max_words is not None and used + para.word_count > max_words:
            break
        page = f" [p. {para.page}]" if para.page is not None else ""
        rendered.append(f"{para.paragraph_id}{page}: {para.text}")
        used += para.word_count
    return "\n\n".join(rendered)


def overview_input(book: Book, emphasis: str, max_structural_words: int) -> str:
    chapter_manifest = "\n".join(
        f"- {c.chapter_id}: {c.title} ({c.word_count} words; kind={c.kind})"
        for c in book.chapters
        if c.kind != "exclude"
    )
    front = [c for c in book.chapters if c.kind == "front"]
    ending = [c for c in book.chapters if c.kind == "ending"]

    if not front:
        front = [c for c in book.chapters if c.kind == "body"][:1]
    if not ending:
        ending = [c for c in book.chapters if c.kind == "body"][-1:]

    per_section_limit = max(1000, max_structural_words // max(1, len(front) + len(ending)))
    structural_texts: list[str] = []
    for chapter in front + ending:
        excerpt = render_paragraphs(book, chapter.paragraph_ids, max_words=per_section_limit)
        structural_texts.append(f"\n## {chapter.chapter_id}: {chapter.title}\n{excerpt}")

    return f"""BOOK TITLE: {book.title}
TOTAL INCLUDED SOURCE WORDS: {book.total_words}
DESIRED READING EMPHASIS: {emphasis}

CHAPTER MANIFEST:
{chapter_manifest}

INTRODUCTION / OPENING AND ENDING MATERIAL:
{''.join(structural_texts)}
"""


def split_paragraph_ids_by_words(book: Book, paragraph_ids: Sequence[str], max_words: int) -> list[list[str]]:
    chunks: list[list[str]] = []
    current: list[str] = []
    current_words = 0
    for pid in paragraph_ids:
        wc = book.paragraphs[pid].word_count
        if current and current_words + wc > max_words:
            chunks.append(current)
            current = []
            current_words = 0
        current.append(pid)
        current_words += wc
    if current:
        chunks.append(current)
    return chunks


OVERVIEW_SYSTEM = """You are performing an inspectional first pass for an extractive abridgement of a rights-cleared nonfiction book.
Your output is analytical metadata only; exact original prose will be selected later.

Classify the work into exactly one nonfiction form: argumentative, historical_investigative_narrative,
biography_memoir, technical_explanatory, case_based_policy_business, or mixed.

Using the section list, opening material, and ending material, identify the governing question, provisional thesis,
and broad chronological or logical arc. Assign chapter priorities. This overview is provisional: later chapter-level
analysis must confirm the author's terms, propositions, supporting arguments, evidence, qualifications, and conclusion.
Do not allow opening setup or concluding interpretation to disappear because middle material is more vivid."""

CHAPTER_ANALYSIS_SYSTEM = """You are performing the analytical-reading pass for one chapter chunk of a rights-cleared nonfiction book.
Your output is analytical metadata only. Do not quote or rewrite the source.

Determine what intellectual function this material performs in the book. Identify questions it addresses, terms it
introduces or clarifies, propositions it advances, reasoning or evidence that supports propositions, qualifications or
limitations, dependencies on earlier material, and its contribution to the unity of the book. Do not treat a vivid
example as a core proposition unless the author uses it to advance the central analysis."""

ANALYTICAL_MAP_SYSTEM = """You are synthesizing an Adlerian analytical map for a nonfiction book from a provisional overview and
chapter-level analyses. Your output is analytical metadata only; exact source passages will be selected later.

Represent the minimum original material a serious reader must retain in order to reconstruct the author's analysis.
Create requirements for: the central question; unity or governing thesis; essential terms in the author's usage; major
propositions; the argument, mechanism, or evidence required for each essential proposition; substantive objections,
limitations, or qualifications; and the concluding answer or implications.

Requirements marked must_be_preserved are hard coverage obligations. Every essential major proposition must have at
least one associated supporting_argument or evidence requirement connected through a supports relation. The
minimum_complete_reading_path must give requirement IDs in the order in which a condensed reader needs them. Use
chapter IDs exactly as supplied. Keep the map discriminating: do not make every illustrative detail essential."""

BASE_CHAPTER_SYSTEM = """You are selecting source passages for a full extractive abridgement of a rights-cleared nonfiction book.
Return paragraph identifiers and analytical tags only; never quote, paraphrase, rewrite, merge, or complete source text.
The application retrieves exact original wording after your selection.

The objective is to preserve the reader's ability to reconstruct the author's analysis. Each nominated passage must
state which analytical requirements it covers. Do not nominate a conclusion alone when the source material here contains
reasoning, evidence, definition, qualification, or objection-response material required to understand why the conclusion
is reached.

READABILITY IS A HARD CONSTRAINT. Each selected block must be a self-contained reading passage, normally 200 to
1,200 words. Start where the author introduces a claim, event, definition, mechanism, evidence, or interpretation; do
not start mid-response, mid-anecdote, or with unresolved referents. End after the local point or event is completed.
Prefer several compact complete passages over one long episode. Do not select bibliography, reference material,
interviews lists, photo inserts, captions, publisher notices, or page furniture.

Each block must be assigned one function: setup, definition, claim, mechanism, evidence, representative_episode,
turning_point, consequence, counterargument, interpretation, or conclusion. Populate covers_requirement_ids only with
requirement IDs supplied in the analytical map. Populate establishes_term_ids with essential-term requirement IDs that
the passage defines or makes intelligible. Populate supports_proposition_ids with major-proposition requirement IDs for
which the passage supplies reasoning or evidence. Populate requires_prior_requirement_ids when the passage would be
unclear without an earlier requirement.

This is the candidate stage. Nominate approximately the requested share of the supplied chapter/chunk. A later global
stage enforces analytical coverage, removes redundancy, and controls the final word budget. Every block ID must begin
with the supplied prefix. Every paragraph ID must come from the supplied input."""

SCORING_SYSTEM = """You are ranking candidate quotation blocks for an analytically readable extractive abridgement.
Your output is analytical metadata only; do not reproduce quotation text. Score each supplied block exactly once.
Reward passages that preserve the unity of the book, establish essential terms, express major propositions, provide the
reasoning or evidence required for propositions, preserve substantive limitations or objections, complete the conclusion,
and remain readable when extracted. Also consider chronological or causal necessity for narrative works. Penalize
redundancy, excessive local detail, and unresolved dependence on omitted context. Identify substantial redundancy using
candidate block IDs."""

QC_SYSTEM = """You are performing final analytical quality control on a proposed extractive abridgement.
Do not rewrite quotations and do not invent block IDs or requirement IDs. Determine whether the retained verbatim
passages enable a serious reader to reconstruct the author's analysis.

Specifically assess whether the retained selection preserves: (1) the central problem and governing thesis; (2) the
structure of the author's answer; (3) essential terminology in the author's usage; (4) each essential major proposition;
(5) sufficient reasoning, evidence, or mechanism for each essential proposition; (6) substantive objections,
limitations, or qualifications; and (7) the concluding answer or implications. Do not approve a selection that retains a
major conclusion while omitting the reasoning required to understand it. Recommend removals only for genuine redundancy,
excessive detail, or unreadability, and additions only from the supplied available candidates. A block protected for an
analytical requirement should not be removed unless another retained block preserves that requirement. Transition notes
are audit comments used to inform the subsequent editorial-transition stage."""

TRANSITION_SYSTEM = """You are writing minimal editorial transitions for a condensed reading edition of a rights-cleared
nonfiction book. The retained passages are verbatim source text. Any transition that appears in the edition is visibly
labelled editorial text and must not imitate or be presented as the author's prose.

Write a transition only when omission creates a material continuity problem. A transition may identify that intervening
discussion, examples, or events have been omitted; state the logical movement between retained passages when grounded in
the supplied analytical map; orient the reader across a chronological shift; or briefly recall a term already established.
It must not add facts, claims, evidence, or evaluation not supported by the supplied passages or analytical requirements.
It must not replace omitted argument with a generated summary or quote the source. In minimal mode, omit a transition
whenever the next retained passage is intelligible without one. In guided mode, a brief orientation bridge may be used
where it materially clarifies the analytical reading path.

Return exactly one record for each supplied candidate gap. If no transition is needed, set necessary to false and text to
an empty string. Use only supplied block IDs and requirement IDs."""

TRANSITION_VALIDATION_SYSTEM = """You are validating reader-visible editorial transitions in an extractive condensed
edition. Approve only transitions that are minimal, visibly editorial in function, and fully grounded in the supplied
retained passages and analytical map. A transition should orient the reader across an omission; it must not substitute
AI-authored substantive content for removed source prose.

Reject or shorten a transition if it adds a factual or interpretive claim not supported by the supplied material,
summarizes omitted evidence or argument in substantive detail, evaluates the author, imitates authorial prose, or is
longer than necessary. Return an empty approved_text when it should not appear."""


def mode_guidance(nonfiction_form: str) -> str:
    if nonfiction_form == "historical_investigative_narrative":
        return """This book is historical/investigative narrative. Preserve chronology and causal development.
Select compact passages that establish institutional origins, major program shifts, decisive episodes,
consequences, and retrospective interpretation. A vivid episode is valuable only insofar as it advances the
historical arc. Penalize extended scene-setting and repeated detail once the significance is established."""
    if nonfiction_form == "biography_memoir":
        return """This book is biographical or memoir-based. Preserve life chronology, formative episodes,
relationships that alter decisions, turning points, and retrospective interpretation. Do not over-retain
atmospheric scenes that do not change the subject's trajectory."""
    if nonfiction_form == "technical_explanatory":
        return """This book is technical/explanatory. Preserve definitions, mechanisms, procedures, key evidence,
representative examples needed for understanding, limitations, and conclusions. Deprioritize repeated
illustrations once a mechanism is clear."""
    if nonfiction_form == "case_based_policy_business":
        return """This book is case-based policy/business nonfiction. Preserve the framework, decisions,
representative cases, outcomes, limitations, and implications. Avoid accumulating multiple cases that make
the same point."""
    if nonfiction_form == "argumentative":
        return """This book is argumentative nonfiction. Preserve thesis, core claims, definitions, supporting
evidence, substantive objections, responses, and conclusion. Deprioritize rhetorical illustration."""
    return """This book has a mixed nonfiction form. Preserve its governing arc while balancing argument,
chronology, evidence, mechanisms, and conclusions. Avoid allowing vivid local material to displace the whole."""


# ---------------------------------------------------------------------------
# Pipeline stages
# ---------------------------------------------------------------------------

def create_structural_overview(
    llm: LLM,
    book: Book,
    emphasis: str,
    max_structural_words: int,
    output_dir: Path,
) -> StructuralOverview:
    LOGGER.info("Generating inspectional structural overview.")
    user_input = overview_input(book, emphasis, max_structural_words)
    overview = llm.structured(OVERVIEW_SYSTEM, user_input, StructuralOverview)
    safe_json_dump(overview, output_dir / "structural_overview.json")
    return overview


def create_chapter_analyses(
    llm: LLM,
    book: Book,
    overview: StructuralOverview,
    chapter_chunk_words: int,
    output_dir: Path,
) -> list[ChapterAnalyticalMap]:
    """Read each included chapter chunk for its analytical function before extracting passages."""
    results: list[ChapterAnalyticalMap] = []
    analysis_dir = output_dir / "chapter_analysis"
    for chapter in [c for c in book.chapters if c.kind != "exclude"]:
        chunks = split_paragraph_ids_by_words(book, chapter.paragraph_ids, chapter_chunk_words)
        LOGGER.info("Analysing intellectual structure of %s (%d chunk(s)).", chapter.title, len(chunks))
        for chunk_index, para_ids in enumerate(chunks, start=1):
            analysis_id = f"{chapter.chapter_id}-A{chunk_index:02d}"
            text = f"""PROVISIONAL BOOK OVERVIEW:
{overview.model_dump_json(indent=2)}

CHAPTER CHUNK TO ANALYSE:
Analysis ID: {analysis_id}
Chapter ID: {chapter.chapter_id}
Chapter title: {chapter.title}
Chunk: {chunk_index} of {len(chunks)}

SOURCE PARAGRAPHS:
{render_paragraphs(book, para_ids)}
"""
            result = llm.structured(CHAPTER_ANALYSIS_SYSTEM, text, ChapterAnalyticalMap)
            if result.chapter_id != chapter.chapter_id or result.analysis_id != analysis_id:
                LOGGER.warning("Normalising returned analysis identifiers for %s.", analysis_id)
                result = result.model_copy(update={
                    "analysis_id": analysis_id,
                    "chapter_id": chapter.chapter_id,
                    "chapter_title": chapter.title,
                    "chunk_index": chunk_index,
                })
            results.append(result)
            safe_json_dump(result, analysis_dir / f"{analysis_id}.json")
    safe_json_dump(results, output_dir / "chapter_analyses.json")
    return results


def validate_analytical_map(analytical_map: AnalyticalMap) -> AnalyticalMap:
    """Validate Adlerian completeness and make support for essential propositions mandatory."""
    req_by_id = {req.requirement_id: req for req in analytical_map.requirements}
    if len(req_by_id) != len(analytical_map.requirements):
        raise RuntimeError("Analytical map contains duplicate requirement IDs.")
    unknown_path = [x for x in analytical_map.minimum_complete_reading_path if x not in req_by_id]
    if unknown_path:
        raise RuntimeError(f"Analytical map contains unknown reading-path requirement IDs: {unknown_path}")
    for relation in analytical_map.argument_relations:
        if relation.source_requirement_id not in req_by_id or relation.target_requirement_id not in req_by_id:
            raise RuntimeError(f"Analytical map relation references unknown requirement: {relation.model_dump()}")

    support_sources_by_claim: dict[str, list[str]] = {}
    for relation in analytical_map.argument_relations:
        if relation.relation == "supports":
            support_sources_by_claim.setdefault(relation.target_requirement_id, []).append(relation.source_requirement_id)

    amended: list[AnalyticalRequirement] = []
    essential_proposition_ids = {
        req.requirement_id for req in analytical_map.requirements
        if req.kind == "major_proposition" and req.must_be_preserved
    }
    for proposition_id in essential_proposition_ids:
        sources = support_sources_by_claim.get(proposition_id, [])
        if not sources:
            raise RuntimeError(
                f"Essential proposition {proposition_id} has no mapped supporting_argument or evidence requirement."
            )
        invalid = [sid for sid in sources if req_by_id[sid].kind not in {"supporting_argument", "evidence"}]
        if invalid:
            raise RuntimeError(
                f"Essential proposition {proposition_id} is supported by non-support requirements: {invalid}"
            )

    required_support_ids = {
        source_id for proposition_id in essential_proposition_ids
        for source_id in support_sources_by_claim.get(proposition_id, [])
    }
    for req in analytical_map.requirements:
        if req.requirement_id in required_support_ids and not req.must_be_preserved:
            amended.append(req.model_copy(update={"must_be_preserved": True, "importance": "essential"}))
        else:
            amended.append(req)
    return analytical_map.model_copy(update={"requirements": amended})


def create_analytical_map(
    llm: LLM,
    book: Book,
    overview: StructuralOverview,
    chapter_analyses: list[ChapterAnalyticalMap],
    output_dir: Path,
) -> AnalyticalMap:
    """Synthesize hard analytical preservation requirements from chapter-level reading."""
    chapter_manifest = "\n".join(
        f"- {c.chapter_id}: {c.title} ({c.word_count} words; kind={c.kind})"
        for c in book.chapters if c.kind != "exclude"
    )
    input_text = f"""BOOK TITLE: {book.title}

CHAPTER MANIFEST:
{chapter_manifest}

PROVISIONAL STRUCTURAL OVERVIEW:
{overview.model_dump_json(indent=2)}

CHAPTER-LEVEL ANALYTICAL MAPS:
{json.dumps([x.model_dump() for x in chapter_analyses], indent=2, ensure_ascii=False)}
"""
    LOGGER.info("Synthesizing Adlerian analytical map.")
    analytical_map = llm.structured(ANALYTICAL_MAP_SYSTEM, input_text, AnalyticalMap)
    analytical_map = validate_analytical_map(analytical_map)
    safe_json_dump(analytical_map, output_dir / "analytical_map.json")
    return analytical_map


def candidate_text_budget(chapter: Chapter, candidate_ratio: float) -> int:
    return max(120, int(chapter.word_count * candidate_ratio))


def requirements_for_chapter(analytical_map: AnalyticalMap, chapter_id: str) -> list[AnalyticalRequirement]:
    local = [
        req for req in analytical_map.requirements
        if not req.preferred_chapter_ids or chapter_id in req.preferred_chapter_ids
    ]
    return local or analytical_map.requirements


def select_candidates_for_chapter(
    llm: LLM,
    book: Book,
    chapter: Chapter,
    overview: StructuralOverview,
    analytical_map: AnalyticalMap,
    emphasis: str,
    candidate_ratio: float,
    chapter_chunk_words: int,
    output_dir: Path,
) -> list[SelectedBlock]:
    chunks = split_paragraph_ids_by_words(book, chapter.paragraph_ids, chapter_chunk_words)
    blocks: list[SelectedBlock] = []
    chapter_function_notes: list[str] = []
    omitted_notes: list[str] = []
    local_requirements = requirements_for_chapter(analytical_map, chapter.chapter_id)

    LOGGER.info("Selecting analytically tagged candidate blocks for %s (%d chunk(s)).", chapter.title, len(chunks))
    for chunk_index, para_ids in enumerate(chunks, start=1):
        chunk_words = sum(book.paragraphs[pid].word_count for pid in para_ids)
        requested_words = max(100, int(chunk_words * candidate_ratio))
        prefix = f"{chapter.chapter_id}-K{chunk_index:02d}-B"
        input_text = f"""BOOK STRUCTURAL OVERVIEW:
{overview.model_dump_json(indent=2)}

ANALYTICAL REQUIREMENTS RELEVANT TO THIS CHAPTER:
{json.dumps([x.model_dump() for x in local_requirements], indent=2, ensure_ascii=False)}

MINIMUM COMPLETE READING PATH:
{', '.join(analytical_map.minimum_complete_reading_path)}

READING EMPHASIS:
{emphasis}

NONFICTION-FORM GUIDANCE:
{mode_guidance(overview.nonfiction_form)}

CHAPTER:
{chapter.chapter_id}: {chapter.title}

CHUNK:
{chunk_index} of {len(chunks)}; {chunk_words} words

CANDIDATE EXTRACTION TARGET:
Nominate approximately {requested_words} words of paragraph blocks from this chunk, corresponding to about {candidate_ratio:.0%} of the source chunk.

REQUIRED BLOCK-ID PREFIX:
{prefix}

SOURCE PARAGRAPHS:
{render_paragraphs(book, para_ids)}
"""
        response = llm.structured(BASE_CHAPTER_SYSTEM + "\n\n" + mode_guidance(overview.nonfiction_form), input_text, ChapterCandidateResponse)
        if response.chapter_id != chapter.chapter_id:
            LOGGER.warning("Response chapter ID %s differed from requested %s; using requested ID.", response.chapter_id, chapter.chapter_id)
        chapter_function_notes.append(response.chapter_contribution)
        omitted_notes.extend(response.omitted_material_description)

        para_set = set(para_ids)
        source_order = {pid: ix for ix, pid in enumerate(para_ids)}
        permitted_requirement_ids = {req.requirement_id for req in analytical_map.requirements}
        output_block_counter = 0
        for block in response.candidate_blocks:
            valid_ids = [pid for pid in block.paragraph_ids if pid in para_set]
            if not valid_ids:
                LOGGER.warning("Ignoring block with no valid paragraph IDs: %s", block.block_id)
                continue
            ordered_ids = sorted(set(valid_ids), key=lambda pid: source_order[pid])
            runs: list[list[str]] = []
            current_run: list[str] = []
            previous_position: int | None = None
            for pid in ordered_ids:
                position = source_order[pid]
                if current_run and previous_position is not None and position != previous_position + 1:
                    runs.append(current_run)
                    current_run = []
                current_run.append(pid)
                previous_position = position
            if current_run:
                runs.append(current_run)

            for run_index, contiguous_ids in enumerate(runs, start=1):
                output_block_counter += 1
                exact_text = "\n\n".join(book.paragraphs[pid].text for pid in contiguous_ids)
                reason = block.selection_reason
                if len(runs) > 1:
                    reason = f"{reason} [Split into contiguous source block {run_index} of {len(runs)}.]"
                blocks.append(
                    SelectedBlock(
                        block_id=f"{prefix}{output_block_counter:02d}",
                        chapter_id=chapter.chapter_id,
                        chapter_title=chapter.title,
                        paragraph_ids=contiguous_ids,
                        word_count=sum(book.paragraphs[pid].word_count for pid in contiguous_ids),
                        selection_reason=reason,
                        importance=block.importance,
                        themes=block.themes,
                        redundancy_risk=block.redundancy_risk,
                        block_function=block.block_function,
                        text=exact_text,
                        covers_requirement_ids=[x for x in block.covers_requirement_ids if x in permitted_requirement_ids],
                        establishes_term_ids=[x for x in block.establishes_term_ids if x in permitted_requirement_ids],
                        supports_proposition_ids=[x for x in block.supports_proposition_ids if x in permitted_requirement_ids],
                        requires_prior_requirement_ids=[x for x in block.requires_prior_requirement_ids if x in permitted_requirement_ids],
                    )
                )

    payload = {
        "chapter_id": chapter.chapter_id,
        "chapter_title": chapter.title,
        "chapter_contribution_notes": chapter_function_notes,
        "omitted_material_description": omitted_notes,
        "candidate_blocks": [asdict(b) for b in blocks],
    }
    safe_json_dump(payload, output_dir / "chapter_candidates" / f"{chapter.chapter_id}.json")
    return blocks


def all_candidate_blocks(
    llm: LLM,
    book: Book,
    overview: StructuralOverview,
    analytical_map: AnalyticalMap,
    emphasis: str,
    candidate_ratio: float,
    chapter_chunk_words: int,
    output_dir: Path,
) -> list[SelectedBlock]:
    candidates: list[SelectedBlock] = []
    for chapter in [c for c in book.chapters if c.kind != "exclude"]:
        candidates.extend(
            select_candidates_for_chapter(
                llm=llm,
                book=book,
                chapter=chapter,
                overview=overview,
                analytical_map=analytical_map,
                emphasis=emphasis,
                candidate_ratio=candidate_ratio,
                chapter_chunk_words=chapter_chunk_words,
                output_dir=output_dir,
            )
        )
    return candidates


def block_preview(block: SelectedBlock, preview_words: int = 180) -> str:
    return truncate_words(block.text, preview_words)


def score_candidates(
    llm: LLM,
    candidates: list[SelectedBlock],
    overview: StructuralOverview,
    analytical_map: AnalyticalMap,
    score_batch_size: int,
    output_dir: Path,
) -> list[SelectedBlock]:
    LOGGER.info("Scoring %d candidate blocks against analytical requirements.", len(candidates))
    manifest = "\n".join(
        f"{b.block_id} | {b.chapter_id} | words={b.word_count} | covers={','.join(b.covers_requirement_ids)} | supports={','.join(b.supports_proposition_ids)} | reason={b.selection_reason}"
        for b in candidates
    )
    scores_by_id: dict[str, ScoredBlockResponse] = {}

    for start in range(0, len(candidates), score_batch_size):
        batch = candidates[start : start + score_batch_size]
        supplied = "\n\n".join(
            f"BLOCK {b.block_id}\nChapter: {b.chapter_id}: {b.chapter_title}\nWords: {b.word_count}\n"
            f"Function: {b.block_function}\nCovers: {', '.join(b.covers_requirement_ids)}\nDefines terms: {', '.join(b.establishes_term_ids)}\n"
            f"Supports propositions: {', '.join(b.supports_proposition_ids)}\nRequires prior: {', '.join(b.requires_prior_requirement_ids)}\n"
            f"Initial importance: {b.importance}\nReason: {b.selection_reason}\nThemes: {', '.join(b.themes)}\nSource preview:\n{block_preview(b)}"
            for b in batch
        )
        user_input = f"""STRUCTURAL OVERVIEW:
{overview.model_dump_json(indent=2)}

ANALYTICAL MAP:
{analytical_map.model_dump_json(indent=2)}

COMPLETE CANDIDATE MANIFEST FOR REDUNDANCY REFERENCE:
{truncate_words(manifest, 7000)}

NONFICTION-FORM GUIDANCE:
{mode_guidance(overview.nonfiction_form)}

SCORING PRINCIPLE:
For argumentative and technical books, give greatest weight to unity, propositions, and their support. For historical or
biographical works, also give substantial weight to chronology, causal development, and interpretation. A concise claim
is not analytically sufficient when it lacks the necessary reasoning or evidence.

BLOCKS TO SCORE IN THIS CALL:
{supplied}
"""
        response = llm.structured(SCORING_SYSTEM, user_input, ScoringResponse)
        for score in response.scores:
            scores_by_id[score.block_id] = score

    for block in candidates:
        score = scores_by_id.get(block.block_id)
        if score is None:
            LOGGER.warning("No score returned for %s; assigning neutral score.", block.block_id)
            block.score = 10.0
            continue
        analytical_value = (
            4 * score.unity_value + 3 * score.term_definition_value + 4 * score.proposition_value
            + 4 * score.argument_support_value + 3 * score.objection_or_limitation_value
            + 3 * score.conclusion_value + 2 * score.analytical_dependency_value
        )
        arc_value = (
            2 * score.chronological_or_logical_necessity
            + 2 * score.institutional_causal_or_argumentative_importance
            + score.turning_point_or_conclusion_value + score.explanatory_density
        )
        if overview.nonfiction_form in {"historical_investigative_narrative", "biography_memoir"}:
            analytical_value, arc_value = analytical_value * 0.8, arc_value * 1.4
        block.score = (
            analytical_value + arc_value + score.readability
            - 2 * score.redundancy_penalty - 2 * score.excessive_detail_penalty
            - 2 * score.context_dependence_penalty
        )
        block.redundant_with = score.redundant_with

    safe_json_dump([asdict(b) for b in candidates], output_dir / "scored_candidates.json")
    return candidates


def chapter_priority_map(overview: StructuralOverview) -> dict[str, str]:
    return {item.chapter_id: item.priority.lower() for item in overview.chapter_priorities}


def chapter_word_caps(
    book: Book,
    target_words: int,
    priority: dict[str, str],
    *,
    min_cap: int = 600,
    high_priority_multiplier: float = 1.15,
) -> dict[str, int]:
    """Per-chapter retained-word ceilings proportional to included source length."""
    included = [c for c in book.chapters if c.kind != "exclude" and c.word_count > 0]
    total_words = sum(c.word_count for c in included)
    if total_words <= 0:
        return {}
    caps: dict[str, int] = {}
    for chapter in included:
        share = chapter.word_count / total_words
        cap = max(min_cap, int(target_words * share))
        if priority.get(chapter.chapter_id) == "high":
            cap = int(cap * high_priority_multiplier)
        caps[chapter.chapter_id] = cap
    return caps


def analytical_coverage(selected: list[SelectedBlock], analytical_map: AnalyticalMap) -> dict[str, list[str]]:
    coverage: dict[str, list[str]] = {req.requirement_id: [] for req in analytical_map.requirements}
    for block in selected:
        tags = set(block.covers_requirement_ids + block.establishes_term_ids)
        for requirement_id in tags:
            if requirement_id in coverage:
                coverage[requirement_id].append(block.block_id)
    return coverage


def choose_blocks_under_budget(
    book: Book,
    candidates: list[SelectedBlock],
    overview: StructuralOverview,
    analytical_map: AnalyticalMap,
    target_ratio: float,
    coverage_mode: str = "all",
) -> tuple[list[SelectedBlock], int]:
    """Select readable blocks while protecting analytical completeness before chapter balance."""
    target_words = int(book.total_words * target_ratio)
    priority = chapter_priority_map(overview)
    chapter_caps = chapter_word_caps(book, target_words, priority)
    min_cap = 600
    chosen: list[SelectedBlock] = []
    chosen_ids: set[str] = set()
    chosen_words_by_chapter: dict[str, int] = {}
    used_words = 0

    included_chapters = [c for c in book.chapters if c.kind != "exclude"]
    by_chapter: dict[str, list[SelectedBlock]] = {}
    for block in candidates:
        by_chapter.setdefault(block.chapter_id, []).append(block)

    def priority_multiplier(block: SelectedBlock) -> float:
        return {"high": 1.18, "medium": 1.0, "low": 0.86}.get(priority.get(block.chapter_id, "medium"), 1.0)

    def utility(block: SelectedBlock) -> float:
        continuity_bonus = 1.07 if block.block_function in {"setup", "turning_point", "consequence", "interpretation", "conclusion"} else 1.0
        analytical_bonus = 1.0 + 0.04 * len(set(block.covers_requirement_ids + block.establishes_term_ids + block.supports_proposition_ids))
        return priority_multiplier(block) * continuity_bonus * analytical_bonus * block.score / math.sqrt(max(block.word_count, 1))

    def cap_for(block: SelectedBlock, protected: bool = False) -> int:
        chapter_cap = chapter_caps.get(block.chapter_id, min_cap)
        if protected:
            return max(chapter_cap, chosen_words_by_chapter.get(block.chapter_id, 0) + block.word_count)
        return chapter_cap

    def conflicts(block: SelectedBlock) -> bool:
        return any(r in chosen_ids for r in block.redundant_with)

    def can_add(block: SelectedBlock, *, protected: bool = False, allow_overshoot: bool = False) -> bool:
        if block.block_id in chosen_ids or (not protected and conflicts(block)):
            return False
        budget_factor = 1.12 if protected else (1.03 if allow_overshoot else 1.0)
        if used_words + block.word_count > int(target_words * budget_factor):
            return False
        current = chosen_words_by_chapter.get(block.chapter_id, 0)
        return current + block.word_count <= cap_for(block, protected=protected)

    def add(block: SelectedBlock, protected: bool = False, requirement_id: str | None = None) -> None:
        nonlocal used_words
        block.protected_anchor = protected or block.protected_anchor
        if requirement_id and requirement_id not in block.protected_requirement_ids:
            block.protected_requirement_ids.append(requirement_id)
        chosen.append(block)
        chosen_ids.add(block.block_id)
        used_words += block.word_count
        chosen_words_by_chapter[block.chapter_id] = chosen_words_by_chapter.get(block.chapter_id, 0) + block.word_count

    req_by_id = {req.requirement_id: req for req in analytical_map.requirements}
    required_ids: list[str] = []
    for requirement_id in analytical_map.minimum_complete_reading_path:
        req = req_by_id.get(requirement_id)
        if req and req.must_be_preserved and requirement_id not in required_ids:
            required_ids.append(requirement_id)
    for req in analytical_map.requirements:
        if req.must_be_preserved and req.requirement_id not in required_ids:
            required_ids.append(req.requirement_id)

    # First pass: protect the intellectual architecture, not merely physical chapter coverage.
    for requirement_id in required_ids:
        req = req_by_id[requirement_id]
        available = [b for b in candidates if requirement_id in b.covers_requirement_ids or requirement_id in b.establishes_term_ids]
        # Candidate blocks must explicitly claim the requirement they satisfy; support metadata is
        # retained as an additional audit relation and is not substituted for requirement coverage.
        for candidate in sorted(available, key=lambda b: (utility(b), -b.word_count), reverse=True):
            if candidate.block_id in chosen_ids:
                if requirement_id not in candidate.protected_requirement_ids:
                    candidate.protected_requirement_ids.append(requirement_id)
                break
            if can_add(candidate, protected=True):
                add(candidate, protected=True, requirement_id=requirement_id)
                break
        else:
            LOGGER.warning("No feasible selected passage protected analytical requirement %s: %s", requirement_id, req.description)

    # Second pass: retain chapter continuity only after essential analytical coverage is protected.
    for chapter in included_chapters:
        available = by_chapter.get(chapter.chapter_id, [])
        if not available or any(b.chapter_id == chapter.chapter_id for b in chosen):
            continue
        must_cover = coverage_mode == "all" or (
            coverage_mode == "major" and (
                chapter.kind in {"front", "ending"} or priority.get(chapter.chapter_id, "medium") in {"high", "medium"}
            )
        )
        if not must_cover:
            continue
        for anchor in sorted(available, key=lambda b: (utility(b), -b.word_count), reverse=True):
            if can_add(anchor, protected=True, allow_overshoot=True):
                add(anchor, protected=True)
                break

    # Third pass: allocate residual budget to the highest-value, non-redundant passages.
    ranked = sorted(candidates, key=utility, reverse=True)
    for block in ranked:
        if can_add(block):
            add(block)

    if used_words < int(target_words * 0.94):
        additions = [b for b in ranked if b.block_id not in chosen_ids and not conflicts(b)]
        additions.sort(key=lambda b: (abs((used_words + b.word_count) - target_words), -utility(b)))
        for block in additions:
            if can_add(block, allow_overshoot=True):
                add(block)
            if used_words >= int(target_words * 0.97):
                break

    chapter_order = {c.chapter_id: i for i, c in enumerate(book.chapters)}
    para_order = {pid: p.index for pid, p in book.paragraphs.items()}
    chosen.sort(key=lambda b: (chapter_order.get(b.chapter_id, 10**6), para_order[b.paragraph_ids[0]]))
    return chosen, target_words


def quality_control(
    llm: LLM,
    overview: StructuralOverview,
    analytical_map: AnalyticalMap,
    selected: list[SelectedBlock],
    candidates: list[SelectedBlock],
    target_words: int,
    output_dir: Path,
) -> QualityControlResponse:
    selected_ids = {b.block_id for b in selected}
    retained_words = sum(b.word_count for b in selected)
    coverage = analytical_coverage(selected, analytical_map)
    retained_manifest = "\n\n".join(
        f"RETAINED {b.block_id} | {b.chapter_id}: {b.chapter_title} | words={b.word_count} | score={b.score:.1f}\n"
        f"Covers: {', '.join(b.covers_requirement_ids)} | Defines: {', '.join(b.establishes_term_ids)} | Supports: {', '.join(b.supports_proposition_ids)}\n"
        f"Protected for: {', '.join(b.protected_requirement_ids)}\nReason: {b.selection_reason}\nPreview: {block_preview(b, 120)}"
        for b in selected
    )
    alternatives = sorted([b for b in candidates if b.block_id not in selected_ids], key=lambda x: x.score, reverse=True)[:30]
    alternative_manifest = "\n".join(
        f"AVAILABLE {b.block_id} | {b.chapter_id}: {b.chapter_title} | words={b.word_count} | score={b.score:.1f} | covers={','.join(b.covers_requirement_ids)} | supports={','.join(b.supports_proposition_ids)} | {b.selection_reason}"
        for b in alternatives
    )
    input_text = f"""STRUCTURAL OVERVIEW:
{overview.model_dump_json(indent=2)}

ANALYTICAL MAP:
{analytical_map.model_dump_json(indent=2)}

PROGRAMMATIC REQUIREMENT COVERAGE BY RETAINED BLOCK:
{json.dumps(coverage, indent=2, ensure_ascii=False)}

WORD BUDGET:
Target: {target_words}
Currently retained: {retained_words}

PROPOSED RETAINED BLOCKS:
{truncate_words(retained_manifest, 15000)}

HIGH-SCORING UNSELECTED ALTERNATIVES:
{alternative_manifest}
"""
    review = llm.structured(QC_SYSTEM, input_text, QualityControlResponse)
    safe_json_dump(review, output_dir / "quality_control.json")
    return review


def apply_qc_changes(
    selected: list[SelectedBlock],
    candidates: list[SelectedBlock],
    review: QualityControlResponse,
    target_words: int,
    book: Book,
    overview: StructuralOverview,
) -> list[SelectedBlock]:
    candidate_map = {b.block_id: b for b in candidates}
    selected_by_chapter: dict[str, list[SelectedBlock]] = {}
    for block in selected:
        selected_by_chapter.setdefault(block.chapter_id, []).append(block)

    remove_ids: set[str] = set()
    for block_id in review.remove_block_ids:
        block = candidate_map.get(block_id)
        if block is None or block.protected_requirement_ids:
            continue
        if len(selected_by_chapter.get(block.chapter_id, [])) <= 1:
            continue
        remove_ids.add(block_id)

    revised = [b for b in selected if b.block_id not in remove_ids]
    current_ids = {b.block_id for b in revised}
    words_by_chapter: dict[str, int] = {}
    for block in revised:
        words_by_chapter[block.chapter_id] = words_by_chapter.get(block.chapter_id, 0) + block.word_count

    priority = chapter_priority_map(overview)
    chapter_caps = chapter_word_caps(book, target_words, priority)
    qc_tolerance = 1.15
    for block_id in review.add_block_ids:
        block = candidate_map.get(block_id)
        if block is None or block_id in current_ids:
            continue
        new_words = sum(b.word_count for b in revised) + block.word_count
        new_chapter_words = words_by_chapter.get(block.chapter_id, 0) + block.word_count
        chapter_cap = int(chapter_caps.get(block.chapter_id, 600) * qc_tolerance)
        if new_words <= int(target_words * 1.12) and (new_chapter_words <= chapter_cap or block.protected_requirement_ids):
            revised.append(block)
            current_ids.add(block_id)
            words_by_chapter[block.chapter_id] = new_chapter_words
    return revised


# ---------------------------------------------------------------------------
# Editorial transition generation and validation
# ---------------------------------------------------------------------------

def selected_in_reading_order(book: Book, selected: list[SelectedBlock]) -> list[SelectedBlock]:
    chapter_order = {chapter.chapter_id: index for index, chapter in enumerate(book.chapters)}
    return sorted(
        selected,
        key=lambda block: (chapter_order.get(block.chapter_id, 10**6), book.paragraphs[block.paragraph_ids[0]].index),
    )


def has_omission_gap(book: Book, first: SelectedBlock, second: SelectedBlock) -> bool:
    if first.chapter_id != second.chapter_id:
        return True
    last_index = book.paragraphs[first.paragraph_ids[-1]].index
    next_index = book.paragraphs[second.paragraph_ids[0]].index
    return next_index > last_index + 1


def potential_transition_pairs(book: Book, selected: list[SelectedBlock]) -> list[tuple[SelectedBlock, SelectedBlock]]:
    ordered = selected_in_reading_order(book, selected)
    return [
        (first, second) for first, second in zip(ordered, ordered[1:])
        if has_omission_gap(book, first, second)
    ]


def _transition_pair_text(
    first: SelectedBlock,
    second: SelectedBlock,
    max_passage_words: int = 450,
) -> str:
    return f"""GAP FROM {first.block_id} TO {second.block_id}
Preceding chapter: {first.chapter_title}
Preceding analytical coverage: {', '.join(first.covers_requirement_ids + first.establishes_term_ids)}
Preceding retained passage excerpt:
{truncate_words(first.text, max_passage_words)}

Following chapter: {second.chapter_title}
Following analytical coverage: {', '.join(second.covers_requirement_ids + second.establishes_term_ids)}
Following passage requires: {', '.join(second.requires_prior_requirement_ids)}
Following retained passage excerpt:
{truncate_words(second.text, max_passage_words)}"""


def generate_editorial_transitions(
    llm: LLM,
    book: Book,
    analytical_map: AnalyticalMap,
    selected: list[SelectedBlock],
    transition_mode: str,
    max_transition_words: int,
    batch_size: int,
    output_dir: Path,
) -> list[EditorialTransition]:
    """Generate disclosed bridges only across actual gaps in the final retained sequence."""
    if transition_mode == "none":
        safe_json_dump([], output_dir / "editorial_transitions.json")
        return []

    pairs = potential_transition_pairs(book, selected)
    if not pairs:
        safe_json_dump([], output_dir / "editorial_transitions.json")
        return []

    requirement_ids = {req.requirement_id for req in analytical_map.requirements}
    selected_ids = {block.block_id for block in selected}
    pair_keys = {(first.block_id, second.block_id) for first, second in pairs}
    transitions: list[EditorialTransition] = []
    LOGGER.info("Generating %s editorial transitions for %d candidate gap(s).", transition_mode, len(pairs))

    for start in range(0, len(pairs), batch_size):
        batch = pairs[start:start + batch_size]
        supplied_pairs = "\n\n---\n\n".join(_transition_pair_text(first, second) for first, second in batch)
        input_text = f"""TRANSITION MODE: {transition_mode}
MAXIMUM WORDS PER TRANSITION: {max_transition_words}

BOOK CENTRAL PROBLEM:
{analytical_map.central_problem}

BOOK UNITY:
{analytical_map.unity_statement}

ANALYTICAL REQUIREMENTS:
{json.dumps([req.model_dump() for req in analytical_map.requirements], indent=2, ensure_ascii=False)}

CANDIDATE GAPS BETWEEN FINAL RETAINED BLOCKS:
{supplied_pairs}
"""
        response = llm.structured(TRANSITION_SYSTEM, input_text, EditorialTransitionResponse)
        for transition in response.transitions:
            key = (transition.before_block_id, transition.after_block_id)
            if key not in pair_keys:
                LOGGER.warning("Ignoring transition for an unknown or non-adjacent block pair: %s -> %s", *key)
                continue
            if not transition.necessary or not transition.text.strip():
                continue
            text = re.sub(r"^\s*(editorial transition|transition)\s*:\s*", "", transition.text.strip(), flags=re.IGNORECASE)
            if word_count(text) > max_transition_words:
                LOGGER.warning("Deferring overlong transition to validation: %s -> %s", *key)
            grounding_blocks = [bid for bid in transition.grounding_block_ids if bid in {key[0], key[1]}]
            grounding_reqs = [rid for rid in transition.grounding_requirement_ids if rid in requirement_ids]
            transitions.append(transition.model_copy(update={
                "text": text,
                "grounding_block_ids": grounding_blocks,
                "grounding_requirement_ids": grounding_reqs,
            }))

    deduplicated: dict[tuple[str, str], EditorialTransition] = {}
    for transition in transitions:
        deduplicated[(transition.before_block_id, transition.after_block_id)] = transition
    result = list(deduplicated.values())
    safe_json_dump(result, output_dir / "editorial_transitions.json")
    return result


def validate_editorial_transitions(
    llm: LLM,
    book: Book,
    analytical_map: AnalyticalMap,
    selected: list[SelectedBlock],
    transitions: list[EditorialTransition],
    max_transition_words: int,
    batch_size: int,
    output_dir: Path,
) -> list[EditorialTransition]:
    """Validate generated bridges before they become reader-visible editorial apparatus."""
    if not transitions:
        safe_json_dump([], output_dir / "editorial_transition_validation.json")
        return []

    ordered = selected_in_reading_order(book, selected)
    block_map = {block.block_id: block for block in ordered}
    transition_map = {(item.before_block_id, item.after_block_id): item for item in transitions}
    validated: list[EditorialTransition] = []
    audits: list[TransitionValidation] = []
    items = list(transition_map.items())
    LOGGER.info("Validating %d proposed editorial transition(s).", len(items))
    for start in range(0, len(items), batch_size):
        batch = items[start:start + batch_size]
        supplied: list[str] = []
        for (before_id, after_id), transition in batch:
            first = block_map[before_id]
            second = block_map[after_id]
            supplied.append(
                _transition_pair_text(first, second)
                + f"\nPROPOSED EDITORIAL TRANSITION: {transition.text}\n"
                + f"DECLARED GROUNDING REQUIREMENTS: {', '.join(transition.grounding_requirement_ids)}"
            )
        transition_separator = "\n\n---\n\n"
        input_text = f"""MAXIMUM APPROVED WORDS PER TRANSITION: {max_transition_words}

BOOK UNITY:
{analytical_map.unity_statement}

ANALYTICAL MAP:
{analytical_map.model_dump_json(indent=2)}

TRANSITIONS TO VALIDATE:
{transition_separator.join(supplied)}
"""
        response = llm.structured(TRANSITION_VALIDATION_SYSTEM, input_text, TransitionValidationResponse)
        audits.extend(response.validations)

    audits_by_pair = {(item.before_block_id, item.after_block_id): item for item in audits}
    for key, transition in transition_map.items():
        check = audits_by_pair.get(key)
        if check is None or not check.grounded or check.introduces_new_claim or check.impersonates_author or check.overly_explanatory:
            LOGGER.info("Omitting unapproved editorial transition %s -> %s.", *key)
            continue
        approved_text = re.sub(
            r"^\s*(editorial transition|transition)\s*:\s*", "", check.approved_text.strip(), flags=re.IGNORECASE
        )
        if not approved_text or word_count(approved_text) > max_transition_words:
            LOGGER.info("Omitting empty or overlong approved editorial transition %s -> %s.", *key)
            continue
        validated.append(transition.model_copy(update={"text": approved_text}))
    safe_json_dump(audits, output_dir / "editorial_transition_validation.json")
    safe_json_dump(validated, output_dir / "editorial_transitions.json")
    return validated


def transition_word_count(transitions: list[EditorialTransition]) -> int:
    return sum(word_count(transition.text) for transition in transitions if transition.necessary)


def enforce_transition_budget(
    book: Book,
    selected: list[SelectedBlock],
    transitions: list[EditorialTransition],
    max_transition_share: float,
    output_dir: Path,
) -> list[EditorialTransition]:
    """Limit generated editorial text without altering approved wording."""
    if not transitions:
        return []
    source_words = sum(block.word_count for block in selected)
    budget_words = int(source_words * max_transition_share)
    ordered_blocks = selected_in_reading_order(book, selected)
    position = {block.block_id: index for index, block in enumerate(ordered_blocks)}
    ordered_transitions = sorted(transitions, key=lambda item: position.get(item.after_block_id, 10**9))
    kept: list[EditorialTransition] = []
    used_words = 0
    for transition in ordered_transitions:
        proposed_words = word_count(transition.text)
        if used_words + proposed_words <= budget_words:
            kept.append(transition)
            used_words += proposed_words
        else:
            LOGGER.info(
                "Omitting approved editorial transition %s -> %s to respect the editorial-word cap (%d words).",
                transition.before_block_id, transition.after_block_id, budget_words,
            )
    safe_json_dump(kept, output_dir / "editorial_transitions.json")
    return kept


# ---------------------------------------------------------------------------
# Output assembly
# ---------------------------------------------------------------------------

def paragraph_range(block: SelectedBlock, book: Book) -> str:
    first = book.paragraphs[block.paragraph_ids[0]]
    last = book.paragraphs[block.paragraph_ids[-1]]
    page_marker = ""
    if first.page is not None:
        page_marker = f", p. {first.page}" if last.page == first.page else f", pp. {first.page}-{last.page}"
    return f"{block.chapter_id}, {block.paragraph_ids[0]}-{block.paragraph_ids[-1]}{page_marker}"


def assemble_reading_markdown(
    book: Book,
    selected: list[SelectedBlock],
    transitions: list[EditorialTransition],
    output_dir: Path,
) -> Path:
    """Create the reading edition from verbatim passages with disclosed editorial bridges."""
    ordered = selected_in_reading_order(book, selected)
    by_chapter: dict[str, list[SelectedBlock]] = {}
    for block in ordered:
        by_chapter.setdefault(block.chapter_id, []).append(block)
    transition_before_block = {
        transition.after_block_id: transition.text.strip()
        for transition in transitions
        if transition.necessary and transition.text.strip()
    }
    first_block_id = ordered[0].block_id if ordered else None

    lines: list[str] = [
        f"# {book.title}",
        "",
        "## Condensed reading edition",
        "",
        "*This edition consists of selected verbatim passages from the original work. Omissions are marked by "
        "centered dots. Brief italicized editorial transitions may be inserted to preserve continuity between "
        "retained passages; these transitions are not part of the original text.*",
        "",
    ]
    for chapter in book.chapters:
        blocks = by_chapter.get(chapter.chapter_id, [])
        if not blocks:
            continue
        lines.extend([f"## {chapter.title}", ""])
        for index, block in enumerate(blocks):
            if block.block_id != first_block_id:
                lines.extend(["", "* * *", ""])
                transition_text = transition_before_block.get(block.block_id)
                if transition_text:
                    lines.extend([f"*Editorial transition: {transition_text}*", ""])
            lines.append(block.text.strip())
            lines.append("")
    path = output_dir / "reading_abridgement.md"
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return path

def assemble_audit_markdown(
    book: Book,
    overview: StructuralOverview,
    analytical_map: AnalyticalMap,
    selected: list[SelectedBlock],
    target_words: int,
    review: QualityControlResponse,
    transitions: list[EditorialTransition],
    output_dir: Path,
) -> Path:
    selected_words = sum(b.word_count for b in selected)
    editorial_words = transition_word_count(transitions)
    coverage = analytical_coverage(selected, analytical_map)
    lines = [
        f"# Selection audit: {book.title}", "",
        f"- Source words analysed: {book.total_words:,}",
        f"- Target words: {target_words:,}",
        f"- Selected verbatim words: {selected_words:,}",
        f"- Retained-source proportion: {selected_words / max(book.total_words, 1):.1%}",
        f"- Editorial transition words: {editorial_words:,}",
        f"- Total reading-edition words: {selected_words + editorial_words:,}", "",
        "## Editorial overview", "", overview.overview, "",
        f"**Central question:** {overview.central_question}", "",
        f"**Governing thesis:** {overview.governing_thesis}", "",
        f"**Nonfiction form:** {overview.nonfiction_form}", "",
        "## Analytical coverage", "",
        "| Requirement | Kind | Importance | Required | Retained block(s) | Status |",
        "|---|---|---|---|---|---|",
    ]
    for req in analytical_map.requirements:
        blocks = coverage.get(req.requirement_id, [])
        status = "covered" if blocks else ("MISSING" if req.must_be_preserved else "not retained")
        desc = req.description.replace("|", "\\|")
        lines.append(
            f"| {req.requirement_id}: {desc} | {req.kind} | {req.importance} | "
            f"{'yes' if req.must_be_preserved else 'no'} | {', '.join(blocks) or '—'} | {status} |"
        )
    lines.extend([
        "", "## Selection balance by chapter", "",
        "| Chapter | Source words | Retained words | Retained share of chapter | Share of abridgement | Status |",
        "|---|---:|---:|---:|---:|---|",
    ])
    retained_by_chapter: dict[str, int] = {}
    for block in selected:
        retained_by_chapter[block.chapter_id] = retained_by_chapter.get(block.chapter_id, 0) + block.word_count
    for chapter in book.chapters:
        if chapter.kind == "exclude":
            continue
        retained = retained_by_chapter.get(chapter.chapter_id, 0)
        status = "represented" if retained else "omitted"
        escaped_title = chapter.title.replace("|", "\\|")
        lines.append(
            f"| {escaped_title} | {chapter.word_count:,} | {retained:,} | "
            f"{retained / max(chapter.word_count, 1):.1%} | {retained / max(selected_words, 1):.1%} | {status} |"
        )
    lines.extend(["", "## Selected passages", ""])
    for block in selected:
        lines.extend([
            f"### {block.block_id}: {block.chapter_title}", "",
            f"- Location: {paragraph_range(block, book)}", f"- Words: {block.word_count}",
            f"- Score: {block.score:.1f}", f"- Function: {block.block_function}",
            f"- Covers analytical requirements: {', '.join(block.covers_requirement_ids) or 'none tagged'}",
            f"- Defines essential terms: {', '.join(block.establishes_term_ids) or 'none tagged'}",
            f"- Supports propositions: {', '.join(block.supports_proposition_ids) or 'none tagged'}",
            f"- Protected requirements: {', '.join(block.protected_requirement_ids) or 'none'}",
            f"- Protected continuity anchor: {block.protected_anchor}", f"- Reason: {block.selection_reason}", "",
        ])
    lines.extend(["## Analytical quality-control assessment", "", review.assessment, ""])
    if review.analytical_coverage.missing_requirement_ids:
        lines.append(f"- Missing required coverage: {', '.join(review.analytical_coverage.missing_requirement_ids)}")
    if review.analytical_coverage.unsupported_proposition_ids:
        lines.append(f"- Unsupported propositions: {', '.join(review.analytical_coverage.unsupported_proposition_ids)}")
    if review.analytical_coverage.undefined_term_ids:
        lines.append(f"- Undefined essential terms: {', '.join(review.analytical_coverage.undefined_term_ids)}")
    if transitions:
        lines.extend(["", "## Approved editorial transitions", ""])
        for transition in transitions:
            lines.extend([
                f"### {transition.before_block_id} to {transition.after_block_id}", "",
                f"- Type: {transition.transition_type}",
                f"- Grounding blocks: {', '.join(transition.grounding_block_ids) or 'adjacent retained passages'}",
                f"- Grounding requirements: {', '.join(transition.grounding_requirement_ids) or 'none tagged'}",
                f"- Text: *{transition.text}*", "",
            ])
    if review.transition_notes:
        lines.extend(["", "## Continuity warnings from quality control", ""])
        for note in review.transition_notes:
            lines.append(f"- Before {note.before_block_id}: {note.note}")
    path = output_dir / "selection_audit.md"
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return path


def assemble_analytical_reading_guide(
    book: Book,
    analytical_map: AnalyticalMap,
    selected: list[SelectedBlock],
    output_dir: Path,
) -> Path:
    """Write a separate analytical guide; never interleave model metadata with source prose."""
    coverage = analytical_coverage(selected, analytical_map)
    block_map = {block.block_id: block for block in selected}
    lines = [
        f"# Analytical reading guide: {book.title}", "",
        "This guide maps the retained verbatim passages to the book's analytical structure. "
        "It is separate from the condensed reading edition, whose substantive content is selected original prose; any generated continuity bridges in the edition are explicitly labelled editorial transitions.", "",
        "## What the book is trying to answer", "", analytical_map.central_problem, "",
        "## Unity of the book", "", analytical_map.unity_statement, "",
        "## Minimum complete reading path", "",
    ]
    req_map = {req.requirement_id: req for req in analytical_map.requirements}
    for index, req_id in enumerate(analytical_map.minimum_complete_reading_path, start=1):
        req = req_map.get(req_id)
        if req is None:
            continue
        blocks = coverage.get(req_id, [])
        locations = "; ".join(
            f"{bid} ({paragraph_range(block_map[bid], book)})" for bid in blocks if bid in block_map
        ) or "Not retained"
        lines.append(f"{index}. **{req.kind.replace('_', ' ').title()} — {req_id}.** {req.description}  ")
        lines.append(f"   Retained passage: {locations}")
    lines.extend(["", "## Essential terms", ""])
    for req in analytical_map.requirements:
        if req.kind == "essential_term":
            lines.append(f"- **{req.requirement_id}.** {req.description} Retained in: {', '.join(coverage.get(req.requirement_id, [])) or 'not retained'}")
    lines.extend(["", "## Major propositions and their support", ""])
    relations_by_target: dict[str, list[str]] = {}
    for relation in analytical_map.argument_relations:
        if relation.relation == "supports":
            relations_by_target.setdefault(relation.target_requirement_id, []).append(relation.source_requirement_id)
    for req in analytical_map.requirements:
        if req.kind != "major_proposition":
            continue
        support_ids = relations_by_target.get(req.requirement_id, [])
        lines.extend([f"### {req.requirement_id}", "", req.description, "", f"Retained proposition passage(s): {', '.join(coverage.get(req.requirement_id, [])) or 'not retained'}", ""])
        if support_ids:
            for support_id in support_ids:
                support = req_map.get(support_id)
                description = support.description if support else support_id
                lines.append(f"- Support {support_id}: {description} Retained in: {', '.join(coverage.get(support_id, [])) or 'not retained'}")
            lines.append("")
    path = output_dir / "analytical_reading_guide.md"
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return path


def _reader_page_size(name: str) -> tuple[float, float]:
    from reportlab.lib.units import inch
    presets = {
        "small-tablet": (7.0 * inch, 10.0 * inch),
        "a5": (5.827 * inch, 8.268 * inch),
        "large-tablet": (7.5 * inch, 10.5 * inch),
    }
    return presets[name]


def _register_reader_fonts(preference: str = "auto") -> dict[str, str]:
    """
    Embed a locally installed serif family when available.
    No font files are bundled with or distributed in this package.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fallback = {
        "regular": "Times-Roman",
        "bold": "Times-Bold",
        "italic": "Times-Italic",
        "bolditalic": "Times-BoldItalic",
        "label": "Times",
    }
    families: list[tuple[str, dict[str, list[str]]]] = [
        (
            "Georgia",
            {
                "regular": [
                    "/System/Library/Fonts/Supplemental/Georgia.ttf",
                    "/Library/Fonts/Georgia.ttf",
                    r"C:\Windows\Fonts\georgia.ttf",
                ],
                "bold": [
                    "/System/Library/Fonts/Supplemental/Georgia Bold.ttf",
                    "/Library/Fonts/Georgia Bold.ttf",
                    r"C:\Windows\Fonts\georgiab.ttf",
                ],
                "italic": [
                    "/System/Library/Fonts/Supplemental/Georgia Italic.ttf",
                    "/Library/Fonts/Georgia Italic.ttf",
                    r"C:\Windows\Fonts\georgiai.ttf",
                ],
                "bolditalic": [
                    "/System/Library/Fonts/Supplemental/Georgia Bold Italic.ttf",
                    "/Library/Fonts/Georgia Bold Italic.ttf",
                    r"C:\Windows\Fonts\georgiaz.ttf",
                ],
            },
        ),
        (
            "DejaVu Serif",
            {
                "regular": [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
                    "/Library/Fonts/DejaVuSerif.ttf",
                ],
                "bold": [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
                    "/Library/Fonts/DejaVuSerif-Bold.ttf",
                ],
                "italic": [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf",
                    "/Library/Fonts/DejaVuSerif-Italic.ttf",
                ],
                "bolditalic": [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSerif-BoldItalic.ttf",
                    "/Library/Fonts/DejaVuSerif-BoldItalic.ttf",
                ],
            },
        ),
    ]
    requested = preference.strip().lower()
    if requested == "times":
        return fallback

    for family_name, candidates in families:
        family_keys = {family_name.lower(), family_name.lower().replace(" ", "")}
        if requested != "auto" and requested not in family_keys:
            continue
        resolved: dict[str, str] = {}
        for style, locations in candidates.items():
            selected_path = next((loc for loc in locations if Path(loc).exists()), None)
            if not selected_path:
                resolved = {}
                break
            resolved[style] = selected_path
        if not resolved:
            continue
        try:
            prefix = family_name.replace(" ", "") + "_Reader"
            registered: dict[str, str] = {}
            for style, selected_path in resolved.items():
                registered_name = f"{prefix}_{style}"
                pdfmetrics.registerFont(TTFont(registered_name, selected_path))
                registered[style] = registered_name
            registered["label"] = family_name
            return registered
        except Exception as exc:
            LOGGER.warning("Unable to embed %s in PDF output: %s. Using built-in Times.", family_name, exc)
    return fallback


def _pdf_safe_text(value: str) -> str:
    from xml.sax.saxutils import escape
    cleaned = normalize_text(value).replace("\u2028", " ").replace("\u2029", " ")
    return escape(cleaned)


def export_reading_pdf(
    book: Book,
    selected: list[SelectedBlock],
    transitions: list[EditorialTransition],
    output_path: Path,
    page_preset: str = "small-tablet",
    body_font_size: float = 14.0,
    font_preference: str = "auto",
) -> None:
    """Render a large-type tablet reading edition directly from selected source blocks."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.pdfbase.pdfmetrics import stringWidth
        from reportlab.platypus import (
            BaseDocTemplate,
            Flowable,
            Frame,
            PageBreak,
            PageTemplate,
            Paragraph,
            Spacer,
        )
    except ImportError as exc:
        raise ImportError(
            "PDF output requires ReportLab. Install dependencies with: pip install -r requirements.txt"
        ) from exc

    page_size = _reader_page_size(page_preset)
    fonts = _register_reader_fonts(font_preference)
    selected_words = sum(block.word_count for block in selected)
    editorial_words = transition_word_count(transitions)

    class OmissionRule(Flowable):
        def __init__(self, width: float = 72.0) -> None:
            super().__init__()
            self.width = width
            self.height = 26.0

        def wrap(self, available_width: float, available_height: float) -> tuple[float, float]:
            self.available_width = available_width
            return available_width, self.height

        def draw(self) -> None:
            from reportlab.lib import colors
            x = (self.available_width - self.width) / 2
            y = self.height / 2
            self.canv.setStrokeColor(colors.HexColor("#D0C8BC"))
            self.canv.setLineWidth(0.55)
            self.canv.line(x, y, x + 22, y)
            self.canv.line(x + self.width - 22, y, x + self.width, y)
            self.canv.setFillColor(colors.HexColor("#82786D"))
            self.canv.setFont(fonts["regular"], 9)
            self.canv.drawCentredString(x + self.width / 2, y - 3.2, "\u2022 \u2022 \u2022")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReaderTitle", parent=styles["Title"], fontName=fonts["bold"], fontSize=27,
        leading=34, alignment=TA_CENTER, textColor=colors.HexColor("#24201C"), spaceAfter=12,
    )
    subtitle_style = ParagraphStyle(
        "ReaderSubtitle", parent=styles["Normal"], fontName=fonts["italic"], fontSize=13,
        leading=20, alignment=TA_CENTER, textColor=colors.HexColor("#60574D"), spaceAfter=8,
    )
    note_style = ParagraphStyle(
        "ReaderNote", parent=styles["Normal"], fontName=fonts["regular"], fontSize=10,
        leading=15, alignment=TA_CENTER, textColor=colors.HexColor("#756C63"),
    )
    chapter_style = ParagraphStyle(
        "ReaderChapter", parent=styles["Heading1"], fontName=fonts["bold"], fontSize=21,
        leading=29, alignment=TA_LEFT, textColor=colors.HexColor("#28231F"), spaceAfter=25,
        keepWithNext=False,
    )
    body_style = ParagraphStyle(
        "ReaderBody", parent=styles["BodyText"], fontName=fonts["regular"], fontSize=body_font_size,
        leading=body_font_size * 1.58, alignment=TA_LEFT, textColor=colors.HexColor("#211E1A"),
        spaceAfter=body_font_size * 0.66, firstLineIndent=body_font_size * 1.42,
        splitLongWords=False, allowWidows=0, allowOrphans=0,
    )
    first_body_style = ParagraphStyle("ReaderFirstBody", parent=body_style, firstLineIndent=0)
    transition_style = ParagraphStyle(
        "ReaderTransition", parent=styles["BodyText"], fontName=fonts["italic"], fontSize=max(10.5, body_font_size - 1.3),
        leading=max(16.0, body_font_size * 1.42), alignment=TA_LEFT, textColor=colors.HexColor("#5E554C"),
        leftIndent=body_font_size * 0.7, rightIndent=body_font_size * 0.7, spaceBefore=5, spaceAfter=14,
        firstLineIndent=0,
    )

    left_margin = 0.67 * inch
    right_margin = 0.67 * inch
    top_margin = 0.72 * inch
    bottom_margin = 0.66 * inch
    frame = Frame(
        left_margin, bottom_margin,
        page_size[0] - left_margin - right_margin,
        page_size[1] - top_margin - bottom_margin,
        leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0, id="reader-frame",
    )
    doc = BaseDocTemplate(
        str(output_path), pagesize=page_size,
        title=f"{book.title} - Condensed Reading Edition",
        leftMargin=left_margin, rightMargin=right_margin, topMargin=top_margin, bottomMargin=bottom_margin,
    )

    def page_decor(canvas: Any, document: Any) -> None:
        page = canvas.getPageNumber()
        if page == 1:
            return
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#756C62"))
        canvas.setStrokeColor(colors.HexColor("#D8D1C7"))
        canvas.setLineWidth(0.45)
        canvas.setFont(fonts["regular"], 8.3)
        header = book.title.upper()
        max_width = page_size[0] - left_margin - right_margin - 16
        while stringWidth(header, fonts["regular"], 8.3) > max_width and len(header) > 20:
            header = header[:-2].rstrip() + "\u2026"
        canvas.drawString(left_margin, page_size[1] - 0.40 * inch, header)
        canvas.line(left_margin, page_size[1] - 0.50 * inch, page_size[0] - right_margin, page_size[1] - 0.50 * inch)
        canvas.setFont(fonts["regular"], 9)
        canvas.drawCentredString(page_size[0] / 2, 0.38 * inch, str(page))
        canvas.restoreState()

    doc.addPageTemplates([PageTemplate(id="reader", frames=[frame], onPage=page_decor)])

    ordered = selected_in_reading_order(book, selected)
    by_chapter: dict[str, list[SelectedBlock]] = {}
    for block in ordered:
        by_chapter.setdefault(block.chapter_id, []).append(block)
    transition_before_block = {
        transition.after_block_id: transition.text.strip()
        for transition in transitions
        if transition.necessary and transition.text.strip()
    }
    first_block_id = ordered[0].block_id if ordered else None

    story: list[Any] = [
        Spacer(1, 1.44 * inch),
        Paragraph(_pdf_safe_text(book.title), title_style),
        Spacer(1, 0.10 * inch),
        Paragraph("Condensed Reading Edition", subtitle_style),
        Spacer(1, 0.40 * inch),
        Paragraph(
            f"Selected verbatim passages from the original work. Omissions are marked discreetly; "
            f"italicized editorial transitions, where present, are not original text. "
            f"Retained source text: {selected_words:,} words; editorial transitions: {editorial_words:,} words.",
            note_style,
        ),
        PageBreak(),
    ]

    first_section = True
    for chapter in book.chapters:
        blocks = by_chapter.get(chapter.chapter_id, [])
        if not blocks:
            continue
        if not first_section:
            story.append(PageBreak())
        first_section = False
        story.extend([Spacer(1, 0.20 * inch), Paragraph(_pdf_safe_text(chapter.title), chapter_style)])
        first_para = True
        for block_index, block in enumerate(blocks):
            if block.block_id != first_block_id:
                story.extend([Spacer(1, 4), OmissionRule(), Spacer(1, 8)])
                transition_text = transition_before_block.get(block.block_id)
                if transition_text:
                    story.append(Paragraph(
                        _pdf_safe_text(f"Editorial transition: {transition_text}"), transition_style
                    ))
                first_para = True
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", block.text.strip()) if p.strip()]
            for paragraph in paragraphs:
                story.append(Paragraph(_pdf_safe_text(paragraph), first_body_style if first_para else body_style))
                first_para = False

    doc.build(story)
    LOGGER.info(
        "Generated reading PDF: %s (%s, %.1f pt body type, %s).",
        output_path, page_preset, body_font_size, fonts["label"],
    )


def export_reading_docx(markdown_path: Path, output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        LOGGER.warning("python-docx is not installed; skipping DOCX export.")
        return

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for line in markdown_path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.strip() == "* * *":
            para = doc.add_paragraph("* * *")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.strip().startswith("*") and line.strip().endswith("*"):
            para = doc.add_paragraph(line.strip().strip("*"))
            para.italic = True
        elif line.strip():
            doc.add_paragraph(line.strip())
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Persistence and execution
# ---------------------------------------------------------------------------

def persist_parsed_book(book: Book, output_dir: Path) -> None:
    metadata = {
        "title": book.title,
        "source_path": book.source_path,
        "total_words": book.total_words,
        "chapters": [asdict(c) for c in book.chapters],
        "diagnostics": book.diagnostics,
    }
    safe_json_dump(metadata, output_dir / "book_metadata.json")
    with (output_dir / "book_paragraphs.jsonl").open("w", encoding="utf-8") as fh:
        for chapter in book.chapters:
            for pid in chapter.paragraph_ids:
                fh.write(json.dumps(asdict(book.paragraphs[pid]), ensure_ascii=False) + "\n")




def write_structure_report(book: Book, output_dir: Path) -> Path:
    excluded_words = sum(c.word_count for c in book.chapters if c.kind == "exclude")
    diagnostics = book.diagnostics or {}
    warnings = diagnostics.get("warnings", [])
    lines = [
        f"# Parsed structure report: {book.title}", "",
        "## Parser confidence", "",
        f"- Format: {diagnostics.get('format', Path(book.source_path).suffix.upper().lstrip('.'))}",
        f"- Navigation source: {diagnostics.get('navigation_source', 'not reported')}",
        f"- Parser confidence: {diagnostics.get('parser_confidence', 'not reported')}",
        f"- Proceed with LLM extraction: {'yes' if diagnostics.get('proceed_with_llm_extraction', True) else 'no'}",
        f"- Fixed-layout indicators: {', '.join(diagnostics.get('fixed_layout_indicators', [])) or 'none detected'}",
        f"- Image-only content risk: {diagnostics.get('image_only_content_risk', 'not assessed')}",
        f"- Unassigned text words: {diagnostics.get('unassigned_text_words', 0):,}",
        f"- Prose fallback items: {diagnostics.get('prose_fallback_items', 0):,}",
        "",
        "## Recovered structure", "",
        f"- Included source words: {book.total_words:,}",
        f"- Excluded/non-reading words: {excluded_words:,}",
        f"- Sections detected: {len(book.chapters)}", "",
        "| ID | Type | Section title | Words |",
        "|---|---|---|---:|",
    ]
    for chapter in book.chapters:
        safe_title = chapter.title.replace("|", "\\|")
        lines.append(f"| {chapter.chapter_id} | {chapter.kind} | {safe_title} | {chapter.word_count:,} |")
    if warnings:
        lines.extend(["", "## Parser warnings", ""])
        lines.extend([f"- {warning}" for warning in warnings])
    lines.extend([
        "",
        "Review this structure before incurring API calls when parser confidence is medium or low. "
        "The program automatically stops before API calls when recovery is insufficient.",
        "",
    ])
    path = output_dir / "parsed_structure_report.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def reset_generated_outputs(output_dir: Path) -> None:
    """Remove prior generated artefacts while leaving unrelated user files untouched."""
    generated_files = [
        "book_metadata.json", "book_paragraphs.jsonl", "parsed_structure_report.md",
        "structural_overview.json", "chapter_analyses.json", "analytical_map.json", "scored_candidates.json", "global_selection.json",
        "quality_control.json", "editorial_transitions.json", "editorial_transition_validation.json",
        "selection_audit.md", "analytical_reading_guide.md", "reading_abridgement.md",
        "reading_abridgement.docx", "reading_abridgement.pdf",
    ]
    for filename in generated_files:
        target = output_dir / filename
        if target.exists():
            target.unlink()
    for generated_dir in ("chapter_candidates", "chapter_analysis"):
        target_dir = output_dir / generated_dir
        if target_dir.exists():
            shutil.rmtree(target_dir)


def run_pipeline(args: argparse.Namespace) -> Path:
    source = Path(args.input_file).expanduser().resolve()
    if not source.exists():
        raise FileNotFoundError(source)
    output_dir = allocate_output_dir(
        Path(args.output_dir),
        source,
        reuse=args.reuse_output_dir,
    )
    if args.reuse_output_dir:
        reset_generated_outputs(output_dir)
    LOGGER.info("Writing outputs to: %s", output_dir)

    LOGGER.info("Loading source book: %s", source)
    book = load_book(source, chapter_map_path=Path(args.chapter_map).expanduser().resolve() if args.chapter_map else None)
    persist_parsed_book(book, output_dir)
    structure_report = write_structure_report(book, output_dir)
    validate_book_structure(book)
    LOGGER.info("Parsed %d chapters/sections and %s included words.", len(book.chapters), f"{book.total_words:,}")
    if args.parse_only:
        LOGGER.info("Parse-only mode complete: %s", structure_report)
        return structure_report

    llm = LLM(model=args.model, retries=args.retries)
    overview = create_structural_overview(
        llm=llm,
        book=book,
        emphasis=args.emphasis,
        max_structural_words=args.max_structural_words,
        output_dir=output_dir,
    )
    chapter_analyses = create_chapter_analyses(
        llm=llm,
        book=book,
        overview=overview,
        chapter_chunk_words=args.chapter_chunk_words,
        output_dir=output_dir,
    )
    analytical_map = create_analytical_map(
        llm=llm,
        book=book,
        overview=overview,
        chapter_analyses=chapter_analyses,
        output_dir=output_dir,
    )

    candidates = all_candidate_blocks(
        llm=llm,
        book=book,
        overview=overview,
        analytical_map=analytical_map,
        emphasis=args.emphasis,
        candidate_ratio=args.candidate_ratio,
        chapter_chunk_words=args.chapter_chunk_words,
        output_dir=output_dir,
    )
    if not candidates:
        raise RuntimeError("No candidate blocks were selected. Review source parsing and model output.")

    scored = score_candidates(
        llm=llm,
        candidates=candidates,
        overview=overview,
        analytical_map=analytical_map,
        score_batch_size=args.score_batch_size,
        output_dir=output_dir,
    )

    selected, target_words = choose_blocks_under_budget(
        book=book,
        candidates=scored,
        overview=overview,
        analytical_map=analytical_map,
        target_ratio=args.target_ratio,
        coverage_mode=args.coverage_mode,
    )
    chapter_caps = chapter_word_caps(book, target_words, chapter_priority_map(overview))
    LOGGER.info(
        "Initial global selection retains %s words against a target of %s.",
        f"{sum(b.word_count for b in selected):,}",
        f"{target_words:,}",
    )

    review = quality_control(
        llm=llm,
        overview=overview,
        analytical_map=analytical_map,
        selected=selected,
        candidates=scored,
        target_words=target_words,
        output_dir=output_dir,
    )
    if args.apply_qc:
        selected = apply_qc_changes(selected, scored, review, target_words, book, overview)

    transitions = generate_editorial_transitions(
        llm=llm,
        book=book,
        analytical_map=analytical_map,
        selected=selected,
        transition_mode=args.transitions,
        max_transition_words=args.max_transition_words,
        batch_size=args.transition_batch_size,
        output_dir=output_dir,
    )
    transitions = validate_editorial_transitions(
        llm=llm,
        book=book,
        analytical_map=analytical_map,
        selected=selected,
        transitions=transitions,
        max_transition_words=args.max_transition_words,
        batch_size=args.transition_batch_size,
        output_dir=output_dir,
    )
    transitions = enforce_transition_budget(
        book=book,
        selected=selected,
        transitions=transitions,
        max_transition_share=args.max_transition_share,
        output_dir=output_dir,
    )

    global_payload = {
        "target_ratio": args.target_ratio,
        "target_words": target_words,
        "selected_words": sum(b.word_count for b in selected),
        "selected_ratio": sum(b.word_count for b in selected) / max(book.total_words, 1),
        "editorial_transition_mode": args.transitions,
        "max_editorial_transition_share": args.max_transition_share,
        "editorial_transition_words": transition_word_count(transitions),
        "total_reading_edition_words": sum(b.word_count for b in selected) + transition_word_count(transitions),
        "editorial_transitions": [transition.model_dump() for transition in transitions],
        "selected_blocks": [asdict(b) for b in selected],
        "quality_control_applied": bool(args.apply_qc),
        "nonfiction_form": overview.nonfiction_form,
        "analytical_requirements": [x.model_dump() for x in analytical_map.requirements],
        "analytical_coverage": analytical_coverage(selected, analytical_map),
        "missing_mandatory_requirement_ids": [
            req.requirement_id for req in analytical_map.requirements
            if req.must_be_preserved and not analytical_coverage(selected, analytical_map).get(req.requirement_id)
        ],
        "coverage_mode": args.coverage_mode,
        "chapter_word_caps": chapter_caps,
    }
    safe_json_dump(global_payload, output_dir / "global_selection.json")

    reading_path = assemble_reading_markdown(book=book, selected=selected, transitions=transitions, output_dir=output_dir)
    assemble_audit_markdown(book=book, overview=overview, analytical_map=analytical_map, selected=selected, target_words=target_words, review=review, transitions=transitions, output_dir=output_dir)
    assemble_analytical_reading_guide(book=book, analytical_map=analytical_map, selected=selected, output_dir=output_dir)
    pdf_path = output_dir / "reading_abridgement.pdf"
    export_reading_pdf(
        book=book,
        selected=selected,
        transitions=transitions,
        output_path=pdf_path,
        page_preset=args.pdf_page_size,
        body_font_size=args.pdf_font_size,
        font_preference=args.pdf_font,
    )
    if not args.no_docx:
        export_reading_docx(reading_path, output_dir / "reading_abridgement.docx")

    LOGGER.info("Completed readable extractive abridgement PDF: %s", pdf_path)
    return pdf_path


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Create a rights-cleared, quotation-dominant extractive abridgement of a nonfiction book using the OpenAI API."
    )
    parser.add_argument("input_file", help="Input EPUB, PDF, DOCX, TXT, or Markdown book file.")
    parser.add_argument(
        "--output-dir",
        default="abridgement_output",
        help="Parent directory for run outputs. By default, each run creates a unique subfolder here.",
    )
    parser.add_argument(
        "--reuse-output-dir",
        action="store_true",
        help="Write directly into --output-dir and replace prior generated artefacts in that folder.",
    )
    parser.add_argument("--chapter-map", help="Optional JSON chapter map for PDFs without reliable embedded bookmarks. Pages are 1-indexed.")
    parser.add_argument("--parse-only", action="store_true", help="Parse and clean the book, write a structure report, and stop before API calls.")
    parser.add_argument(
        "--model",
        default=os.environ.get("OPENAI_MODEL", "gpt-5-mini"),
        help="OpenAI model supporting structured outputs. Defaults to OPENAI_MODEL or gpt-5-mini.",
    )
    parser.add_argument("--target-ratio", type=float, default=0.25, help="Final retained source-word ratio. Default: 0.25.")
    parser.add_argument(
        "--coverage-mode",
        choices=["all", "major", "none"],
        default="all",
        help="Continuity coverage after mandatory analytical coverage. 'all' attempts to retain at least one passage per parsed non-back-matter section; default: all.",
    )
    parser.add_argument(
        "--candidate-ratio",
        type=float,
        default=0.42,
        help="Candidate quotation share requested per chapter before global pruning. Default: 0.42.",
    )
    parser.add_argument(
        "--emphasis",
        default="Preserve the governing thesis, major claims, essential definitions and concepts, causal or technical mechanisms, decisive evidence, necessary case material, important counterarguments or limitations, and conclusions or implications.",
        help="Selection emphasis supplied to the model.",
    )
    parser.add_argument(
        "--chapter-chunk-words",
        type=int,
        default=18000,
        help="Split unusually long chapters above this word count for selection calls. Default: 18000.",
    )
    parser.add_argument(
        "--max-structural-words",
        type=int,
        default=24000,
        help="Maximum introduction/ending words supplied to the overview call. Default: 24000.",
    )
    parser.add_argument("--score-batch-size", type=int, default=20, help="Candidate blocks per scoring call. Default: 20.")
    parser.add_argument("--retries", type=int, default=3, help="Maximum retries per API call. Default: 3.")
    parser.add_argument(
        "--apply-qc",
        action="store_true",
        help="Apply final quality-control add/remove recommendations when they remain within the word budget tolerance.",
    )
    parser.add_argument(
        "--transitions",
        choices=["none", "minimal", "guided"],
        default="minimal",
        help="Reader-visible editorial transition policy. Default: minimal; transitions are visibly labelled and are not source text.",
    )
    parser.add_argument(
        "--max-transition-words",
        type=int,
        default=45,
        help="Maximum words permitted in an approved editorial transition. Default: 45.",
    )
    parser.add_argument(
        "--transition-batch-size",
        type=int,
        default=12,
        help="Candidate discontinuities per transition generation or validation call. Default: 12.",
    )
    parser.add_argument(
        "--max-transition-share",
        type=float,
        default=0.02,
        help="Maximum editorial-transition words as a proportion of retained verbatim words. Default: 0.02.",
    )
    parser.add_argument(
        "--pdf-page-size",
        choices=["small-tablet", "a5", "large-tablet"],
        default="small-tablet",
        help="PDF page preset. Default: small-tablet (7 x 10 inches).",
    )
    parser.add_argument(
        "--pdf-font-size",
        type=float,
        default=14.0,
        help="Body type size in points for the PDF reading edition. Default: 14.0.",
    )
    parser.add_argument(
        "--pdf-font",
        choices=["auto", "georgia", "dejavu serif", "dejavuserif", "times"],
        default="auto",
        help="Serif font for the PDF. Auto embeds Georgia or DejaVu Serif when available; otherwise uses Times.",
    )
    parser.add_argument("--no-docx", action="store_true", help="Do not generate an optional DOCX copy of the abridgement.")
    parser.add_argument("--verbose", action="store_true", help="Enable detailed logging.")
    return parser


def validate_args(args: argparse.Namespace) -> None:
    if not 0 < args.target_ratio < 1:
        raise ValueError("--target-ratio must be between 0 and 1.")
    if not args.target_ratio < args.candidate_ratio <= 1:
        raise ValueError("--candidate-ratio must be greater than --target-ratio and no greater than 1.")
    if args.chapter_chunk_words < 1000:
        raise ValueError("--chapter-chunk-words must be at least 1000.")
    if args.score_batch_size < 1:
        raise ValueError("--score-batch-size must be positive.")
    if not 10 <= args.max_transition_words <= 100:
        raise ValueError("--max-transition-words must be between 10 and 100.")
    if args.transition_batch_size < 1:
        raise ValueError("--transition-batch-size must be positive.")
    if not 0.0 <= args.max_transition_share <= 0.10:
        raise ValueError("--max-transition-share must be between 0 and 0.10.")
    if not 11.0 <= args.pdf_font_size <= 20.0:
        raise ValueError("--pdf-font-size must be between 11 and 20 points.")


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
    )
    try:
        validate_args(args)
        output = run_pipeline(args)
    except Exception as exc:
        LOGGER.exception("Pipeline failed: %s", exc)
        return 1
    print(f"\nReadable abridgement written to: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
